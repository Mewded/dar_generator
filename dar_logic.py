import os
import re
from datetime import datetime

import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_dar_summary(input_pdf, output_folder):
    summary_text = extract_summary(input_pdf)  # your parsing logic here
    
    # create output filename
    output_path = os.path.join(output_folder, "DAR_Report_Output.pdf")
    
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    story = [Paragraph("Daily Activity Report Summary", styles["Title"]), Spacer(1, 12)]
    story.append(Paragraph(summary_text, styles["Normal"]))
    doc.build(story)
    
    return output_path


# -------- CONFIG (same-folder I/O) --------
BASE_DIR   = os.getcwd()
INPUT_FILE = os.path.join(BASE_DIR, "DAR Report.pdf")   # raw logbook export
LOGO_FILE  = os.path.join(BASE_DIR, "logo.png")     # optional logo
# -----------------------------------------

# Target sections & order
SECTIONS = [
    "Incident Reports (IR) / Alarms",
    "Elevator Entrapment Incidents",
    "SPD Presence/Emergency Response on Site",
    "Property Damage",
    "Tenant Issues",
    "Retail Issues",
    "Transient Removal",
    "Key Service (Lock & Unlock)",
    "Loading Dock Access (Lock & Unlock)",
    "Fire Panel Bypass/Online",
    "AES Phone Calls",
    "Work Orders",
    "Janitorial",
    "Additional Information",
]

# Verb normalization to past tense
VERB_MAP = {
    "open": "opened",
    "unlock": "unlocked",
    "lock": "locked",
    "secure": "secured",
    "issue": "issued",
    "collect": "collected",
    "deliver": "delivered",
    "remove": "removed",
    "escort": "escorted",
    "close": "closed",
    "receive": "received",
    "call": "called",
    "handle": "handled",
    "extend": "extended",
    "bypass": "bypassed",
    "bring": "brought",
    "brought": "brought",   # already past
    # ✅ irregulars / common typos
    "put": "put",
    "was": "was",
    "putted": "put",
    "puted": "put",
    "set": "set",
    "hit": "hit",
    "cut": "cut",
    "shut": "shut",
    "leave": "left",
}

DATETIME_RX = re.compile(r"\b(\d{1,2}/\d{1,2}/\d{4})\s+(\d{1,2}:\d{2}\s*[AP]M)\b", re.IGNORECASE)

def to_past_tense(text: str) -> str:
    if not text:
        return text
    words = text.split()
    if not words:
        return text
    first = words[0].lower()
    if first in VERB_MAP:
        words[0] = VERB_MAP[first]
    elif not first.endswith("ed"):
        words[0] = first + ("d" if first.endswith("e") else "ed")
    return " ".join(words)

def extract_text_lines(file_path):
    lines = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            lines.extend(ln.strip() for ln in t.splitlines())
    return lines

def parse_date_range(lines, input_file):
    """
    1) 'Period <start> - <end>'
    2) 'Date Range <start> - <end>'
    3) Fallback: min..max of all timestamps found anywhere (incl. TOUR)
    Returns: header_text, filename_token
    """
    raw_text = "\n".join(lines)

    m = re.search(
        r"Period\s*:?[\s\n]*([0-9/]+\s+\d{1,2}:\d{2}\s*[AP]M)\s*[-–]\s*([0-9/]+\s+\d{1,2}:\d{2}\s*[AP]M)",
        raw_text, re.IGNORECASE
    )
    if not m:
        m = re.search(
            r"Date\s*Range\s*:?[\s\n]*([0-9/]+\s+\d{1,2}:\d{2}\s*[AP]M)\s*[-–]\s*([0-9/]+\s+\d{1,2}:\d{2}\s*[AP]M)",
            raw_text, re.IGNORECASE
        )
    if m:
        start, end = m.group(1).strip(), m.group(2).strip()
        start_simple = start.split()[0].replace("/", "-")
        end_simple   = end.split()[0].replace("/", "-")
        return f"{start} – {end}", f"{start_simple}_to_{end_simple}"

    # Fallback: scan all timestamps in the doc
    all_matches = DATETIME_RX.findall(raw_text)
    if all_matches:
        dts = []
        for d, t in all_matches:
            try:
                dts.append(datetime.strptime(f"{d} {t.upper()}", "%m/%d/%Y %I:%M %p"))
            except Exception:
                pass
        if dts:
            sdt = min(dts)
            edt = max(dts)
            header = f"{sdt.strftime('%m/%d/%Y %I:%M %p')} – {edt.strftime('%m/%d/%Y %I:%M %p')}"
            token  = f"{sdt.strftime('%m-%d-%Y')}_to_{edt.strftime('%m-%d-%Y')}"
            return header, token

    # Last resort: filename like "09-18-25 to 09-19-25"
    fn = os.path.basename(input_file)
    mf = re.search(r"(\d{2}\s*-\s*\d{2}\s*-\s*\d{2})\s*to\s*(\d{2}\s*-\s*\d{2}\s*-\s*\d{2})", fn, re.IGNORECASE)
    if mf:
        s = mf.group(1).replace(" ", "")
        e = mf.group(2).replace(" ", "")
        return f"{s} to {e}", f"{s}_to_{e}"

    return "Unknown Date Range", "unknown"

def bold_officer(name: str) -> str:
    """
    Normalize officer names to 'Officer First Last' format with bold styling.
    Handles cases like 'Faiz Mohmand' (already correct) or 'ALI Kassim' (last first).
    """
    n = (name or "").strip()
    if not n:
        return ""

    # Remove any leading 'Officer' or 'S/O'
    n = re.sub(r"^(Officer|S/O)\s+", "", n, flags=re.IGNORECASE)

    parts = n.split()
    if len(parts) == 2:
        first, second = parts
        # If the first part is ALL CAPS and the second looks like Firstname -> assume it's "LAST FIRST"
        if first.isupper() and second[0].isupper() and second[1:].islower():
            n = f"{second.capitalize()} {first.capitalize()}"
        else:
            n = f"{first.capitalize()} {second.capitalize()}"
    else:
        # Capitalize each part normally
        n = " ".join(p.capitalize() for p in parts)

    return f"<b>Officer {n}</b>"


def classify(buffer, labels):
    """
    Decide which summary section this event belongs to.
    Uses explicit category lines first, then heuristics.
    """
    # 1) Explicit category wins
    if "category" in buffer:
        cat = buffer["category"].lower()
        if "aes phone call" in cat:
            return "AES Phone Calls"
        if ("loading dock" in cat or "dock gate" in cat) and "abm" not in buffer.get("action", "").lower():
            return "Loading Dock Access (Lock & Unlock)"
        if "key service" in cat or "key" in cat:
            return "Key Service (Lock & Unlock)"
        if "bypass online" in cat or "fire panel" in cat or "fire system online" in cat  or "until" in cat:
            return "Fire Panel Bypass/Online"
        if "transient" in cat:
            return "Transient Removal"
        if "work order" in cat:
            return "Work Orders"
        if "retail" in cat:
            return "Retail Issues"
        if "incident report" in cat or "alarm" in cat:
            return "Incident Reports (IR) / Alarms"
        if "janitorial" in cat:
            return "Janitorial"
        if "other/miscellaneous" in cat:
            return "Additional Information"
 

    # 2) Full-text heuristics
    txt = " ".join([
        *labels,
        buffer.get("action", ""),
        buffer.get("company", ""),
        buffer.get("location", ""),
        buffer.get("category", "")
    ]).lower()

    # 🧭 Elevator Entrapment — handle first to prevent other rules from overriding
    if any(k in txt for k in [
        "elevator entrapment incident",
        "stuck in elevator",
        "elevator incident",
        "got stuck in cap",
        "doors stayed closed",
        "kone technician",
        "otis elevator"
    ]):
        return "Elevator Entrapment Incidents"
    
    # Tenant Issues — context-based, exclude elevator cases
    if (
        "tenant" in txt
        and any(k in txt for k in [
            "issue", "concern", "complaint", "problem",
            "request", "notify", "notified", "reported"
        ])
        and not any(n in txt for n in [
            "elevator", "entrapment", "stuck in elevator", "kone", "otis"
        ])
    ):
        return "Tenant Issues"

    # 🔧 Property Damage (check BEFORE generic loading dock or IR)
    if any(w in txt for w in [
        "damage", "damaged", "bent", "broken", "crack", "dent",
        "unable to close", "hit", "struck", "collision", "impact"
    ]):
        if any(n in txt for n in [
            "gate", "door", "frame", "lock", "glass", "loading dock", "dock gate"
        ]):
            return "Property Damage"
    
    # Elevator Entrapment or Stuck Elevator (strict, not Tenant)
    if ("elevator" in txt or "entrapment" in txt or "stuck in elevator" in txt) and "tenant" not in txt:
        return "Elevator Entrapment Incidents"
    
        # 🆕 SPD Presence / Emergency Response on Site — strict detection (same pattern as Elevator)
    if any(k in txt for k in [
        "spd presence/emergency response on site",
        "spd presence",
        "emergency response on site",
        "spd response",
        "sfd medics",
        "911 called",
        "police responded",
        "medical emergency on site",
        "officer contacted spd",
        "security called 911",
        "security called spd"
    ]):
        return "SPD Presence/Emergency Response on Site"
    
    # 🆕 Heuristic fallback if label is messy
    if ("spd" in txt and "response" in txt) or "emergency response on site" in txt:
        return "SPD Presence/Emergency Response on Site"
    
    # Retail / Tenant handling fallback
    if "tenant" in txt and any(k in txt for k in [
        "issue", "concern", "complaint", "problem", "request", "notify", "notified", "reported"
    ]):
        if not any(n in txt for n in ["elevator", "entrapment", "stuck in elevator"]):
            return "Tenant Issues"
        
    if "aes" in txt or "phone call" in txt:
        return "AES Phone Calls"

    # Loading dock movements (but not ABM janitorial tasks)
    if ("loading dock" in txt or "dock gate" in txt) and not any(k in txt for k in ["abm notified", "upload picture"]):
        return "Loading Dock Access (Lock & Unlock)"

    if "key service" in txt or re.search(r"\bkey\s*(lock|unlock|service|issued|return|pickup|drop|set)?\b", txt):
        return "Key Service (Lock & Unlock)"

    if any(k in txt for k in ["panel", "bypass", "trbl", "supv", "fire alarm", "alarm test", "hold"]):
        return "Fire Panel Bypass/Online"

    if "transient" in txt or "trespass" in txt or ("removed" in txt and "person" in txt):
        return "Transient Removal"

    if "work order" in txt or "building engines" in txt:
        return "Work Orders"

    # 🚫 Do NOT treat the generic noun "incident" as IR.
    # Only accept strong IR signals:
    if ("incident report" in txt) or re.search(r"\b(ir|incident)\s*#\s*\d+", txt) or \
       any(k in txt for k in ["police", "911", "injury", "assault", "theft", "robbery"]):
        return "Incident Reports (IR) / Alarms"

    # Stronger Janitorial detection
    if (
        buffer.get("category", "").lower() in ["janitorial", "seattle ambassadors"]
        or "janitorial" in txt
        or "abm notified" in txt
        or "upload picture" in txt
        or "abm" in txt
        or "clean" in txt
        or "trash" in txt
        or "garbage" in txt
        or "spill" in txt
        or "vacuum" in txt
        or "mop" in txt
        or "sweep" in txt
        or "ambassador" in txt        
        or "mid call" in txt          
        or "mid dispatch" in txt     
        or "seattle ambassadors" in txt
    ):
        return "Janitorial"

    return "Unclassified"



def _extract_dt(line: str):
    """
    Try to pull a datetime object from a built event line like:
    '09/18/25 8:18 AM – Officer ...'
    """
    m = re.match(r"^(\d{2}/\d{2}/\d{2})\s+(\d{1,2}:\d{2}\s*[AP]M)", line)
    if not m:
        return datetime.max  # fallback, push to end
    try:
        return datetime.strptime(f"{m.group(1)} {m.group(2)}", "%m/%d/%y %I:%M %p")
    except Exception:
        return datetime.max

def parse_events(lines):
    """
    Walk the document lines and extract NEW ACTIVITY blocks into our sections.
    - Completely ignore TOUR blocks for events.
    - Officer name comes from nested '- Officer :' lines under "Officer",
      or from lines like 'ALI Kassim (Officers)'.
    - Category labels (e.g., 'Loading Dock Gate', 'Key Service', 'AES Phone Call', 'Transient Removal') are captured and used to classify.
    - Multi-line 'action' answers are merged.
    - Break/Lunch/Category noise is skipped.
    """
    parsed = {s: [] for s in SECTIONS}

    in_tour = False
    buffer = {}
    labels = set()
    waiting_officer_nested = False
    last_field = None

    # --- NEW: transient counting support ---
    transient_count = 0
    transient_tag_seen = False
    # ---------------------------------------

    # --- NEW: collect Incident Reports grouped by (date, officer) ---
    incident_groups = {}
    # ---------------------------------------------------------------

    # --- NEW: detect Seattle Ambassadors blocks ---
    in_ambassador = False
    ambassador_buffer = {}
    # ------------------------------------------------

    # --- NEW: detect Unsecure door blocks ---
    in_unsecure = False 
    unsecure_buffer = {}    

    # --- NEW: detect SPD Presence / Emergency Response blocks ---
    in_spd = False
    spd_buffer = {}

    # lines that indicate a new field/segment (not a continuation)
    starters = [
        "Start Date", "Officer", "- Officer :", "Date & Time", "Date/Time", "- Date :", "- Time :",
        "Details", "Call Details", "Location", "- Location :", "Company", "- Company", "Vendor",
        # include the Comments block used by Other/Miscellaneous
        "Comments", "- Multi-line text field",
        "Geolocation", "Evidence", "NEW ACTIVITY", "TOUR", "Start Time", "End Time", "Report Details",
        "Posts included", "Activities included", "New Group", "Tags", "Duration", "Max. Tour Duration",
        "- Picture", "Picture", "Key Service", "Loading Dock Gate", "Fire Panel", "Janitorial",
        "Transient Removal", "Retail Issues", "Tenant Issues", "Fire Panel Bypass/Online", "Incident",
        "Totals Activities", "Total Activities", "Activity Duration", "Object Duration",
        "AES Phone Call", "Work Order", 
        # Label itself may appear as a line
        "Other/Miscellaneous",
        # 👇 ADDED so these headers don't concatenate into narratives
        "Synopsis", "Follow up", "Escalation?", "- Upload picture"
    ]

    # NEW: accept any "… (Officers)" line as the officer source
    OFFICER_LINE_RX = re.compile(r"^(.*?)\s*\(Officers?\)\s*$", re.IGNORECASE)
    # NEW: Also accept "(Site Supervisors)" if the first pattern fails
    OFFICER_LINE_RX_ALT = re.compile(r"^(.*?)\s*\(Site\s+Supervisors?\)\s*$", re.IGNORECASE)

    MULTILINE_RX    = re.compile(r"^-\s*Multi-?line\s+text\s+field\s*:\s*(.*)$", re.IGNORECASE)

    def is_new_block_line(ln: str) -> bool:
        if not ln:
            return True
        for s in starters:
            if ln.startswith(s):
                return True
        if "(Officers)" in ln:
            return True
        # A raw header-style timestamp line
        if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*[AP]M$", ln, re.IGNORECASE):
            # 🚩 Ignore these if we’re inside an Incident Report
            if buffer.get("category", "").lower().startswith("incident report"):
                return False
            return True
        
        if ln in ("300 Pine Street", "300 Pine Street Call Details"):
            return True
        return False

    def _fmt_date_for_line(datestr: str) -> str:
        # Format like your build_event_line does (MM/DD/YY HH:MM AM/PM)
        m = re.search(r"(\d{1,2})/(\d{1,2})/(\d{4})\s+(\d{1,2}:\d{2}\s*[AP]M)", (datestr or ""), re.IGNORECASE)
        if m:
            mm, dd, yyyy, tm = m.groups()
            return f"{int(mm):02d}/{int(dd):02d}/{yyyy[-2:]} {tm.upper()}"
        return datestr or ""

    def flush_event():
        nonlocal buffer, labels, last_field, transient_count, transient_tag_seen, incident_groups

        # 🚫 Skip empty or incomplete buffers (e.g., only officer name with no date/action)
        if not buffer or (
            not buffer.get("action")
            and not buffer.get("incident_description")
            and not buffer.get("incident_comments")
            and not buffer.get("category")
        ):
            buffer.clear()
            last_field = None
            return

        # Determine section
        sec = classify(buffer, labels)

        # Default action for AES when not provided
        if sec == "AES Phone Calls" and not buffer.get("action"):
            buffer["action"] = "handled AES phone call to put the fire system on test"
        
        # Default for Fire Panel if action missing
        # if sec == "Fire Panel Bypass/Online" and not buffer.get("action"):
        #     buffer["action"] = "updated the fire panel status"
        
        # INCIDENT REPORTS: use the same pattern as other sections
        if sec == "Incident Reports (IR) / Alarms":           
            # Force officer line to always use Start Date - use this to se the start date
            if buffer.get("start_date"):
                buffer["date"] = buffer["start_date"]
            else:
                buffer["date"] = buffer.get("date", "")
            # 👉 DO NOT override buffer["date"] anymore — keep Start Date for officer line

            # Merge description + comments into narrative
            desc = (buffer.get("incident_description") or "").strip()
            cmts = (buffer.get("incident_comments") or "").strip()
            parts = []
            if desc:
                parts.append(desc.rstrip(".") + ".")
            if cmts:
                parts.append(cmts.rstrip(".") + ".")
            narrative = " ".join(parts).strip()

            # Add vehicle info if present
            vehicle_bits = []
            if buffer.get("vehicle_description"):
                vehicle_bits.append(buffer["vehicle_description"].strip())
            mm = " ".join([
                (buffer.get("color") or "").strip(),
                (buffer.get("make") or "").strip(),
                (buffer.get("model") or "").strip()
            ]).strip()
            if mm:
                vehicle_bits.append(mm)
            if vehicle_bits:
                narrative += f" Vehicle described as {', '.join(vehicle_bits)}."

            # Add incident-specific info (date + time + location)
            extra_info = []
            incident_date = (buffer.get("incident_date") or "").strip()
            incident_time = (buffer.get("incident_time") or "").strip()
            incident_location = (buffer.get("location") or "").strip()
            
            # 🧭 Smart location inference (only if no explicit location and narrative exists)
            if not incident_location and (desc or cmts):
                source_text = (desc + " " + cmts).lower()

                # 1️⃣ Look for common prepositions like "at", "in", "on", "near"
                m_explicit = re.search(
                    r"\b(?:at|in|on|inside|near|around)\s+([A-Za-z0-9\-\s]+?)(?:[.,;]|$)",
                    source_text,
                    re.I,
                )
                if m_explicit:
                    incident_location = m_explicit.group(1).strip(" .,-")

                # 2️⃣ Guess short codes like SB, NB, P1, L2, Lobby, Roof, etc.
                if not incident_location:
                    m_code = re.search(
                        r"\b([A-Z]{1,3}\d?|L\d|P\d|Dock|Garage|Lobby|Roof|Basement|Floor|Entrance)\b",
                        source_text,
                        re.I,
                    )
                    if m_code:
                        incident_location = m_code.group(1).strip()

                # 3️⃣ Handle floor references
                if not incident_location and "floor" in source_text:
                    m_floor = re.search(r"on\s+the\s+([A-Za-z0-9\s]+?floor)", source_text)
                    if m_floor:
                        incident_location = m_floor.group(1).strip()

                # 4️⃣ Keep only first mention if multiple found
                if incident_location and "," in incident_location:
                    incident_location = incident_location.split(",")[0].strip()

                # 5️⃣ Normalize case
                if incident_location:
                    incident_location = re.sub(r"\s+", " ", incident_location).strip()
                    if re.match(r"^[A-Z]{1,3}\d?$", incident_location):
                        incident_location = incident_location.upper()
                    else:
                        incident_location = incident_location.title()

            # 🧹 Stop location from swallowing text from next sections
            incident_location = re.split(
                r"\b(Synopsis|All persons involved|Who Called|Vehicle Information|Evidence|numbers\))\b",
                incident_location,
                maxsplit=1,
                flags=re.IGNORECASE
            )[0].strip(" ,:-")

            # --- Try to detect embedded "Incident Date" inside narrative if not already parsed ---
            if not buffer.get("incident_date"):
                m_embedded = re.search(r"Incident Date[:\s]+([0-9/]+)\s*(?:at\s+([0-9:]+\s*[APap][Mm]))?", narrative)
                if m_embedded:
                    buffer["incident_date"] = m_embedded.group(1)
                    if m_embedded.group(2):
                        buffer["incident_time"] = m_embedded.group(2).strip()

            # --- Try to detect embedded "Location" if not parsed ---
            if not buffer.get("location"):
                m_loc = re.search(r"Location[:\s]+([A-Za-z0-9#\s\-]+)", narrative)
                if m_loc:
                    buffer["location"] = m_loc.group(1).strip()
            
            # ✅ Normalize & format the location name for consistency
            if buffer.get("location"):
                buffer["location"] = format_location_name(buffer["location"])
                incident_location = buffer["location"]

            if incident_date:
                # --- normalize incident time to AM/PM format ---
                formatted_time = None
                if incident_time:
                    time_clean = incident_time.strip().lower().replace("hrs", "").replace(":", "").replace(" ", "")
                    match_24h = re.match(r"^(\d{1,2})(\d{2})?$", time_clean)
                    if match_24h:
                        hh = int(match_24h.group(1))
                        mm = match_24h.group(2) or "00"
                        ampm = "AM"
                        if hh >= 12:
                            ampm = "PM"
                            if hh > 12:
                                hh -= 12
                        elif hh == 0:
                            hh = 12
                        formatted_time = f"{hh}:{mm} {ampm}"
                    else:
                        # if already has AM/PM
                        formatted_time = incident_time.upper().replace("HRS", "").strip()

                # --- build red-highlighted info ---
                if formatted_time:
                    extra_info.append(
                        f"<font color='red'>Incident Date: <b>{incident_date}</b> at <b>{formatted_time}</b></font>"
                    )
                else:
                    extra_info.append(
                        f"<font color='red'>Incident Date: <b>{incident_date}</b></font>"
                    )

            # 🟥 Use explicit or inferred location; never leave blank if one is found
            final_location = incident_location or buffer.get("location", "").strip()
            if final_location:
                extra_info.append(f"<font color='red'>Location: <b>{final_location}</b></font>")
            else:
                extra_info.append(f"<font color='red'>Location: <b>N/A</b></font>")


            # ✅ Add Who Called (or fallback to officer)
            who_called = (buffer.get("who_called") or "").strip()

            # 🧩 If "Who Called" missing, fallback to officer name (properly formatted)
            if not who_called and buffer.get("officer"):
                who_called = f"Officer {buffer['officer'].strip()}"

            # 🧠 Normalize Who Called only if it looks like "LASTNAME FIRSTNAME"
            if who_called:
                # Remove double spaces and extra parentheses
                who_called = re.sub(r"\s{2,}", " ", who_called).strip()
                who_called = re.sub(r"\(.*?\)", "", who_called).strip()

                # Detect and fix reversed names like "KING Jovonne"
                if re.match(r"^([A-Z]{2,})\s+([A-Z][a-z]+)$", who_called):
                    parts = who_called.split()
                    who_called = f"{parts[1].capitalize()} {parts[0].capitalize()}"

                # Fix case for names that already contain "Officer"
                if who_called.lower().startswith("officer "):
                    name_part = who_called[8:].strip()
                    # handle pattern "KING Jovonne"
                    m = re.match(r"([A-Z]{2,})\s+([A-Z][a-z]+)", name_part)
                    if m:
                        first = m.group(2).capitalize()
                        last = m.group(1).capitalize()
                        who_called = f"Officer {first} {last}"
                    else:
                        who_called = "Officer " + " ".join(w.capitalize() for w in name_part.split())

                extra_info.append(f"<font color='black'>Who Called: <b>{who_called}</b></font>")


            # ✅ Add Parties Involved
            if buffer.get("parties_involved"):
                parties = buffer["parties_involved"].strip()
                parties = re.sub(r"\s{2,}", " ", parties)
                extra_info.append(f"<font color='black'>Parties Involved: <b>{parties}</b></font>")

            extra_text = ""
            if extra_info:
                extra_text = " (" + ", ".join(extra_info) + ")"

            # Save as action (like other sections)
            # 🧠 Smart prefix: lowercase first letter unless it's a name or number
            narrative_clean = narrative.strip()
            if narrative_clean:
                # Lowercase only if the first word isn't a name/acronym (starts with uppercase followed by lowercase)
                if re.match(r"^[A-Z][a-z]", narrative_clean):
                    narrative_clean = narrative_clean[0].lower() + narrative_clean[1:]
                buffer["action"] = f"reported that {narrative_clean}".strip() + extra_text
            else:
                buffer["action"] = "reported that an incident occurred on site." + extra_text


            # 🚩 Suppress duplicate trailing location
            loc = buffer.pop("location", None)
            # --- Final safeguard: ensure timestamp before building event (with AM/PM) ---
            if not buffer.get("date") and buffer.get("incident_date"):
                # normalize to AM/PM if needed
                t = (buffer.get("incident_time") or "").strip()
                if t:
                    t_clean = t.lower().replace("hrs", "").replace(":", "").replace(" ", "")
                    match_24h = re.match(r"^(\d{1,2})(\d{2})?$", t_clean)
                    if match_24h:
                        hh = int(match_24h.group(1))
                        mm = match_24h.group(2) or "00"
                        ampm = "AM"
                        if hh >= 12:
                            ampm = "PM"
                            if hh > 12:
                                hh -= 12
                        elif hh == 0:
                            hh = 12
                        formatted_time = f"{hh}:{mm} {ampm}"
                    else:
                        # already contains AM/PM
                        formatted_time = t.upper().replace("HRS", "").strip()
                    buffer["date"] = f"{buffer['incident_date']} {formatted_time}"
                else:
                    buffer["date"] = buffer["incident_date"]


            # Build and append event line
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)

            # Restore location in case buffer is reused
            if loc:
                buffer["location"] = loc
        # --- END Incident Reports handling ---

        # ELEVATOR ENTRAPMENT: final clean and professional version (IR-style)
        elif sec == "Elevator Entrapment Incidents":
            
            buffer.pop("start_date", None)
            buffer.pop("date", None)
            
            # --- Force officer line to always use Start Date (ignore later overwrites) ---
            if buffer.get("start_date"):
                start_dt = buffer["start_date"].strip()
                # Normalize to short year format and consistent AM/PM spacing
                start_dt = re.sub(r"(\d{4})", lambda m: m.group(1)[-2:], start_dt)
                start_dt = re.sub(r"\s*([APap][Mm])", lambda m: " " + m.group(1).upper(), start_dt)
                buffer["date"] = start_dt
                buffer["timestamp_locked"] = True  # prevent later overwrite
            else:
                buffer["date"] = buffer.get("date", "")

            desc = (buffer.get("incident_description") or "").strip()
            action = (buffer.get("action") or "").strip()
            location = (buffer.get("location") or "").strip()
            company = (buffer.get("company") or "KONE Elevator Company").strip()
            incident_date = (buffer.get("incident_date") or "").strip()

            # 🧭 Smart location inference (only if no explicit location and narrative exists)
            if not location and (desc or action):
                source_text = (desc or action).lower()

                # 1️⃣ Look for common prepositions
                m_explicit = re.search(
                    r"\b(?:at|in|on|inside|near|around)\s+([A-Za-z0-9\-\s]+?)(?:[.,;]|$)",
                    source_text,
                    re.I,
                )
                if m_explicit:
                    location = m_explicit.group(1).strip(" .,-")

                # 2️⃣ Guess short codes like SB, NB, P1, etc.
                if not location:
                    m_code = re.search(
                        r"\b([A-Z]{1,3}\d?|L\d|P\d|Dock|Garage|Lobby|Roof|Basement)\b",
                        source_text,
                        re.I,
                    )
                    if m_code:
                        location = m_code.group(1).strip()

                # 3️⃣ Handle “floor” or area mentions
                if not location and "floor" in source_text:
                    m_floor = re.search(r"on\s+the\s+([A-Za-z0-9\s]+?floor)", source_text)
                    if m_floor:
                        location = m_floor.group(1).strip()

                # 4️⃣ Only keep the first inferred match (avoid multi-location)
                if location and "," in location:
                    location = location.split(",")[0].strip()

                # 5️⃣ Normalize capitalization
                if location:
                    location = re.sub(r"\s+", " ", location).strip()
                    if re.match(r"^[A-Z]{1,3}\d?$", location):
                        location = location.upper()
                    else:
                        location = location.title()

            # Prefer long incident description if available
            narrative = desc if desc else action

            # 🧹 Clean unwanted prefixes and fragments from description
            narrative = re.sub(r"(?i)\b(long\s*)?description\s*of\s*incident\s*:?", "", narrative).strip()
            # narrative = re.sub(r"(?i)^at\s*\d{1,2}:\d{2}\s*(am|pm)?\s*,?\s*", "", narrative).strip()
            narrative = re.sub(r"(?i)-\s*long description of incident.*", "", narrative).strip()
            narrative = re.sub(r"–\s*\(.*long description of incident.*\)", "", narrative, flags=re.IGNORECASE).strip()

            # Detect elevator car identifier (e.g., "FRT elevator 13", "Cap 10", "Car 3")
            elevator_id = ""
            m = re.search(r"(frt\s*elevator\s*\d+|cap\s*\d+|car\s*\d+)", narrative, re.IGNORECASE)
            if m:
                elevator_id = m.group(0).strip()

            # 🧭 Construct the action text
            buffer["action"] = (
                f"responded to an elevator entrapment "
                f"{('in ' + elevator_id) if elevator_id else 'at ' + (location or 'the site')}, "
                f"confirmed the occupant’s safety while awaiting technician arrival, "
                f"and informed {company} for immediate service response and release."
            )

            # If a clear narrative exists, use that instead
            if narrative:
                buffer["action"] = narrative.strip().rstrip(".")

            # 🧠 Add smart narrative prefix: "reported that ..." if not already present
            officer_name = (buffer.get("officer") or "").strip()
            if narrative:
                # lower-case first letter safely
                narrative_clean = narrative[0].lower() + narrative[1:] if len(narrative) > 1 else narrative
                # avoid duplicating prefix if it already says "reported that"
                if not re.search(r"(?i)\breported that\b", narrative_clean):
                    if officer_name:
                        buffer["action"] = f"reported that {narrative_clean}"
                    else:
                        buffer["action"] = f"reported that {narrative_clean}"
            else:
                buffer["action"] = "reported that an incident occurred on site."
            
            # --- 🕒 Final safeguard: ensure timestamp before building event (with AM/PM) ---
            if not buffer.get("date") and buffer.get("incident_date"):
                t = (buffer.get("incident_time") or "").strip()
                if t:
                    # normalize to AM/PM if needed
                    t_clean = t.lower().replace("hrs", "").replace(":", "").replace(" ", "")
                    match_24h = re.match(r"^(\d{1,2})(\d{2})?$", t_clean)
                    if match_24h:
                        hh = int(match_24h.group(1))
                        mm = match_24h.group(2) or "00"
                        ampm = "AM"
                        if hh >= 12:
                            ampm = "PM"
                            if hh > 12:
                                hh -= 12
                        elif hh == 0:
                            hh = 12
                        formatted_time = f"{hh}:{mm} {ampm}"
                    else:
                        # if already has AM/PM
                        formatted_time = t.upper().replace("HRS", "").strip()
                    buffer["date"] = f"{buffer['incident_date']} {formatted_time}"
                else:
                    buffer["date"] = buffer["incident_date"]

            # Build the main event line (includes timestamp + officer)
            evt = build_event_line(buffer)

            if evt:
                # Clean stray fragments
                evt = re.sub(r"\s*-\s*Long Description of Incident\s*:?\s*", " ", evt, flags=re.IGNORECASE).strip()
                evt = re.sub(r"–\s*\(.*long description of incident.*\)", "", evt, flags=re.IGNORECASE).strip()
                evt = re.sub(r"–\s*\(.*?\)$", "", evt, flags=re.DOTALL).strip()

                # 🧹 Clean the location field itself
                location_clean = re.sub(r"(?i)(-?\s*long\s*description\s*of\s*incident\s*:?.*)", "", location).strip()

                # ✅ Format location name for consistency (e.g., Rooftop → Rooftop, fcc → FCC)
                if location_clean:
                    location_clean = format_location_name(location_clean)

                # 🕓 Normalize Incident Time → AM/PM format (same logic as IR)
                incident_time = (buffer.get("incident_time") or "").strip()
                formatted_time = ""
                if incident_time:
                    time_clean = incident_time.lower().replace("hrs", "").replace(":", "").replace(" ", "")
                    match_24h = re.match(r"^(\d{1,2})(\d{2})?$", time_clean)
                    if match_24h:
                        hh = int(match_24h.group(1))
                        mm = match_24h.group(2) or "00"
                        ampm = "AM"
                        if hh >= 12:
                            ampm = "PM"
                            if hh > 12:
                                hh -= 12
                        elif hh == 0:
                            hh = 12
                        formatted_time = f"{hh}:{mm} {ampm}"
                    else:
                        formatted_time = incident_time.upper().replace("HRS", "").strip()

                # 🟥 Build final red-highlighted info (IR-style ordering)
                extra_info = []
                if incident_date and formatted_time:
                    extra_info.append(
                        f"<font color='red'>Incident Date: <b>{incident_date}</b> at <b>{formatted_time}</b></font>"
                    )
                elif incident_date:
                    extra_info.append(
                        f"<font color='red'>Incident Date: <b>{incident_date}</b></font>"
                    )
                else:
                    extra_info.append(
                        f"<font color='red'>Incident Date: <b>N/A</b></font>"
                    )

                # 🟥 Location finalization (use explicit or inferred; never show N/A if we can find one)
                final_location = location_clean or location
                if final_location:
                    extra_info.append(
                        f"<font color='red'>Location: <b>{final_location}</b></font>"
                    )
                else:
                    extra_info.append(
                        f"<font color='red'>Location: <b>N/A</b></font>"
                    )

                # ✅ Add Parties Involved (if available)
                parties_raw = (buffer.get("parties_involved") or "").strip()
                if parties_raw:
                    parties_clean = re.sub(r"(?i)\b(photos?|evidence)\b.*", "", parties_raw).strip()
                    parties_clean = re.sub(r"\s{2,}", " ", parties_clean)
                    extra_info.append(
                        f"<font color='black'>Parties Involved: <b>{parties_clean}</b></font>"
                    )

                # Combine and finalize
                extra_text = " (" + ", ".join(extra_info) + ")"
                evt = evt.rstrip(".") + extra_text

                # Append cleaned result
                parsed[sec].append(evt)
        # --- END Elevator Entrapment handling ---

        # ✅ Enhanced Work Orders handling
        elif sec == "Work Orders":
            # 🧹 Always reset old accidental description from previous section
            buffer.pop("description", None)
            
            # --- Force officer line to always use Start Date (ignore later overwrites) ---
            if buffer.get("start_date"):
                start_dt = buffer["start_date"].strip()
                # Normalize to short year format and consistent AM/PM spacing
                start_dt = re.sub(r"(\d{4})", lambda m: m.group(1)[-2:], start_dt)
                start_dt = re.sub(r"\s*([APap][Mm])", lambda m: " " + m.group(1).upper(), start_dt)
                buffer["date"] = start_dt
                buffer["timestamp_locked"] = True  # marker to prevent later overwrites
            else:
                buffer["date"] = buffer.get("date", "")
            # --- Pull fields ---
            description = (buffer.get("description") or "").strip()
            # --- Normalize lines for Work Orders so we can safely search between headers ---
            normalized_lines = []
            for ln in lines:
                ln = re.sub(r"([A-Za-z])-(\s*[A-Za-z])", r"\1\n-\2", ln)
                ln = re.sub(r"(Order)(Start\s*Date)", r"\1\n\2", ln, flags=re.IGNORECASE)
                normalized_lines.extend(ln.splitlines())

            # --- Extract clean Description ONLY between '- Description :' and '- Work Order Placed...' (after Upload picture) ---
            description = ""
            upload_idx = None
            desc_idx = None
            placed_idx = None

            # Step 1: find "Upload picture", "Description", and "Work Order Placed" lines
            for idx, l in enumerate(normalized_lines):
                if re.search(r"-\s*upload\s*pictures?\s*:", l, re.IGNORECASE):
                    upload_idx = idx
                if upload_idx is not None and re.search(r"-\s*description\s*:", l, re.IGNORECASE):
                    desc_idx = idx
                if re.search(r"-\s*work\s*order\s*placed\s*on\s*building\s*engines", l, re.IGNORECASE):
                    placed_idx = idx
                    break

            # Step 2: collect only lines between description and work order placed
            if desc_idx is not None and placed_idx and placed_idx > desc_idx:
                val_lines = []
                for nxt in normalized_lines[desc_idx:placed_idx]:
                    s = nxt.strip()
                    if not s:
                        continue
                    # Skip headers/junk/footer
                    if re.search(r"^-+\s*description\s*:\s*", s, re.IGNORECASE):
                        # remove the "- Description :" label itself
                        s = re.sub(r"^-+\s*description\s*:\s*", "", s, flags=re.IGNORECASE)
                    if re.search(r"REPORT\s*-\s*LOGBOOK\s*PDF", s, re.IGNORECASE):
                        continue
                    if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}", s):  # footer timestamps
                        continue
                    if re.search(r"\b300\s+Pine\s+Street\b", s, re.IGNORECASE):
                        continue
                    if s:
                        val_lines.append(s)

                # Join and normalize
                description = " ".join(val_lines).strip()
                description = re.sub(r"\s*\.\s*\.?", ".", description)
                description = re.sub(r"\s+", " ", description).strip()

            location    = (buffer.get("location") or "").strip()
            action_raw  = (buffer.get("action") or "").lower()

            # 🧭 Smart location inference (only if no explicit location and description exists)
            if not location and description:
                desc_text = description.lower()

                # 1️⃣ Look for common prepositions (in/on/at/near)
                m_explicit = re.search(
                    r"\b(?:at|in|on|inside|near|around)\s+([A-Za-z0-9\-\s]+?)(?:[.,;]|$)",
                    desc_text,
                    re.I,
                )
                if m_explicit:
                    location = m_explicit.group(1).strip(" .,-")

                # 2️⃣ Guess short codes (SB, NB, L1, P1, etc.)
                if not location:
                    m_code = re.search(
                        r"\b([A-Z]{1,3}\d?|L\d|P\d|Dock|Garage|Lobby|Roof|Basement)\b",
                        description,
                        re.I,
                    )
                    if m_code:
                        location = m_code.group(1).strip()

                # 3️⃣ Handle "floor" or area references
                if not location and "floor" in desc_text:
                    m_floor = re.search(r"on\s+the\s+([A-Za-z0-9\s]+?floor)", desc_text)
                    if m_floor:
                        location = m_floor.group(1).strip()

                # 4️⃣ Normalize capitalization for readability
                if location:
                    location = re.sub(r"\s+", " ", location).strip()
                    if re.match(r"^[A-Z]{1,3}\d?$", location):
                        location = location.upper()
                    else:
                        location = location.title()
            
            # ✅ Format location name for consistency
            if location:
                location = format_location_name(location)

            # --- Detect explicit Building Engines field ---
            placed_on_be = False
            for key, val in buffer.items():
                key_l = str(key).lower()
                val_l = str(val).lower()
                if "work order placed on building engines" in key_l or "work order placed on building engines" in val_l:
                    if "yes" in val_l:
                        placed_on_be = True
                    break

            # --- Skip shift-handover noise / “generic” WO without description ---
            if any(x in action_raw for x in [
                "new emails received",
                "work orders communicated",
                "important info passed",
                "shift",
            ]):
                # do NOT append anything for handover summaries
                pass
            elif not description and not placed_on_be:
                # no description and no BE flag → likely a generic click, skip
                pass
            else:
                # --- Normalize and merge multi-line description ---
                description = re.sub(r"\s*\.\s*\.?", ".", description)
                description = re.sub(r"\s+", " ", description).strip().strip(". ")
                

                # --- Vendor detection (optional polish) ---
                m_vendor = re.search(
                    r"\b(Cedar Grove|ABM|FedEx|UPS|SPS|Ryder|DHL|Old Dominion|USPS|CORT|Corti|Canteen)\b",
                    description, re.IGNORECASE
                )
                vendor_name = m_vendor.group(1) if m_vendor else None

                # --- Build polished narrative ---
                if description:
                    action_text = f"documented a work order indicating that {description[0].lower() + description[1:]}"
                else:
                    action_text = "documented a work order request at the site"

                if vendor_name:
                    action_text += f", and notified {vendor_name} for service"

                if placed_on_be:
                    action_text += ". <font color='green'>Work order placed on Building Engines.</font>"
                else:
                    action_text += ". <font color='red'>Pending submission to Building Engines.</font>"

                # Append location once (here), and prevent second append in build_event_line
                # ✅ Always append location clearly (either explicit or inferred)
                if location:
                    # If "Location:" not already present, add it in consistent format
                    if not re.search(r"\(Location\s*:", action_text, re.I):
                        action_text += f" – <font color='red'>(Location: <b>{location}</b>)</font>"

                buffer["action"]  = action_text.strip()
                buffer["company"] = ""     # avoid duplicate "for Company"
                if "location" in buffer:   # one-location safeguard
                    buffer["location_for_display"] = buffer.pop("location")

                evt = build_event_line(buffer)

                # restore location for safety
                if "location_for_display" in buffer:
                    buffer["location"] = buffer.pop("location_for_display")
    
                # Append once
                if evt:
                    parsed[sec].append(evt)
        # --- END Enhanced Work Orders handling ---

        # TENANT ISSUES: improve professionalism and clarity

        # PROPERTY DAMAGE: operational clarity style
        elif sec == "Property Damage":
            action = (buffer.get("action") or "").strip()
            location = (buffer.get("location") or "").strip()
            company = (buffer.get("company") or "").strip()
            lower = action.lower()

            # --- Keyword detection for type of damage ---
            if any(k in lower for k in ["hit", "struck", "collided", "impact", "crash", "bump"]):
                buffer["action"] = (
                    f"reported property damage at {location or 'the site'} after impact involving "
                    f"{company or 'a vehicle'}, noting {action}"
                )

            elif any(k in lower for k in ["broken", "cracked", "shattered", "smashed", "glass", "window"]):
                buffer["action"] = (
                    f"reported property damage involving glass or structural breakage at {location or 'the site'}, "
                    f"with details indicating {action}"
                )

            elif any(k in lower for k in ["bent", "dented", "warped", "lock", "frame", "gate", "door"]):
                buffer["action"] = (
                    f"reported property damage at {location or 'the site'}, describing physical issues such as {action}"
                )

            elif any(k in lower for k in ["burn", "scorch", "fire", "heat"]):
                buffer["action"] = (
                    f"reported property damage related to fire or heat exposure at {location or 'the site'}, "
                    f"with details noting {action}"
                )

            elif any(k in lower for k in ["leak", "flood", "water", "spill"]):
                buffer["action"] = (
                    f"reported property damage associated with water intrusion at {location or 'the site'}, "
                    f"with details noting {action}"
                )

            else:
                # General fallback
                buffer["action"] = (
                    f"reported property damage at {location or 'the site'}, with details noting {action}"
                )

            # --- Append final event line ---
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)
               
        # KEY SERVICE: make output more professional & contextual
        elif sec == "Key Service (Lock & Unlock)":
            # --- Fetch and pre-clean action ---
            raw_action = (buffer.get("action") or "")
            raw_action = clean_shift_noise(raw_action)

            company = (buffer.get("company") or "").strip()
            location = (buffer.get("location") or "").strip()

            # --- Smart defaults if company missing ---
            if not company:
                if "pine" in location.lower() or "3rd" in location.lower():
                    company = "Victrola Coffee"
                elif "4th" in location.lower() or "uniqlo" in location.lower():
                    company = "UNIQLO"

            # --- Skip if no valid action remains ---
            if not raw_action.strip():
                pass
            else:
                action = to_past_tense(raw_action.strip())

                # --- Remove redundancy and repeated company/door words ---
                if company:
                    # Build flexible pattern: match either full company name OR its first word (e.g., "Victrola" from "Victrola Coffee")
                    first_word = company.split()[0] if company else ""
                    pattern = rf"(\bthe\s+)?\b({re.escape(company)}|{re.escape(first_word)})\b(\s+the\b)?(\s+doors?\b)?"

                    action = re.sub(pattern, "", action, flags=re.IGNORECASE)

                # Remove duplicate "the the" / "doors doors" / trailing "the"
                action = re.sub(r"\bthe\s+the\b", "the", action, flags=re.IGNORECASE)
                action = re.sub(r"\bdoors\s+doors\b", "doors", action, flags=re.IGNORECASE)
                action = re.sub(r"\bthe\s*$", "", action, flags=re.IGNORECASE)
                action = re.sub(r"\s{2,}", " ", action).strip()

                # 🧠 Normalize capitalization globally (e.g., "Secured" → "secured", "Unlocked" → "unlocked")
                if re.match(r"^[A-Z][a-z]+\b", action):
                    first_word = action.split()[0].lower()
                    if first_word in [
                        "secured", "locked", "unlocked", "granted", "provided",
                        "issued", "returned", "escorted", "assisted", "facilitated",
                        "supervised", "verified", "ensured", "conducted", "closed"
                    ]:
                        action = action[0].lower() + action[1:]

                # --- Unlock scenarios ---
                if "unlock" in action.lower() or "gave access" in action.lower() or "give access" in action.lower():
                    # Detect "gave access to X" and optional "for Y"
                    recipient_match = re.search(r"gave\s+access\s+to\s+([A-Za-z\s]+?)(?:\s+for\s+([A-Za-z\s]+))?(?:\s|$)", action, re.IGNORECASE)
                    recipient = ""
                    requester = ""

                    if recipient_match:
                        recipient = recipient_match.group(1).strip()
                        requester = recipient_match.group(2).strip() if recipient_match.group(2) else ""

                    if "delivery" in action.lower():
                        buffer["action"] = (
                            f"conducted key service and unlocked the {company} doors, "
                            f"granting access to delivery personnel for scheduled drop-off"
                        )
                    elif recipient and requester:
                        buffer["action"] = (
                            f"conducted key service and granted access to {recipient} for {requester} "
                            f"through the {company or 'designated area'} doors"
                        )
                    elif recipient:
                        buffer["action"] = (
                            f"conducted key service and granted access to {recipient} "
                            f"through the {company or 'designated area'} doors"
                        )
                    elif "request" in action.lower():
                        delivery_item = ""
                        for word in ["pastry", "supplies", "equipment", "package", "shipment", "delivery"]:
                            if word in action.lower():
                                delivery_item = f" {word}"
                                break
                        buffer["action"] = (
                            f"conducted key service in response to a request from {company or 'delivery personnel'} "
                            f"and unlocked the {company} doors, granting secure access for the scheduled{delivery_item} delivery"
                        )
                    elif "customer" in action.lower():
                        buffer["action"] = (
                            f"conducted key service and unlocked the {company} doors, "
                            f"providing access to customers during business hours"
                        )
                    elif "event" in action.lower() or "contractor" in action.lower():
                        buffer["action"] = (
                            f"conducted key service and unlocked the {company or 'designated area'} doors, "
                            f"facilitating access for event staff or contractors"
                        )
                    else:
                        buffer["action"] = (
                            f"conducted key service and unlocked the {company} doors, "
                            f"ensuring authorized access for scheduled activity"
                        )
                # --- Lock / Secure scenarios ---
                elif any(word in action.lower() for word in ["lock", "secure", "close", "closed", "closing"]):
                    if any(word in action.lower() for word in ["close", "closed", "closing", "end of shift", "finished work"]):
                        buffer["action"] = (
                            f"conducted key service and {action} the {company} doors, "
                            f"securing the premises at the end of operations"
                        )
                    elif "after" in action.lower() or "finished" in action.lower():
                        buffer["action"] = (
                            f"conducted key service and {action} the {company} doors, "
                            f"securing the area after completion of scheduled work"
                        )
                    else:
                        buffer["action"] = (
                            f"conducted key service and {action} the {company} doors, "
                            f"ensuring proper security of the location"
                        )
                
                # --- Handle issuing and returning of keys/badges ---
                elif re.search(r"\b(issued|provided|handed)\b", action, re.IGNORECASE):
                    # Extract possible key/badge identifiers
                    item_list = re.findall(r"\b(key\d*|badge\d*|key|badge|keys|badges)\b", action, re.IGNORECASE)
                    item_list = [i.lower() for i in item_list]

                    # 🧠 Smart article + plural logic
                    def smart_item_phrase(item_text: str) -> str:
                        """
                        Adds natural 'the', 'a', or plural handling for key/badge phrases.
                        Example:
                        key  → 'a key'
                        keys → 'keys'
                        key1 → 'the key1'
                        badge A → 'the badge A'
                        badge → 'a badge'
                        """
                        if not item_text:
                            return ""
                        txt = item_text.strip().lower()
                        # plural -> keep as-is
                        if txt.endswith("s") and not re.search(r"\d", txt):
                            return item_text.strip()
                        # numbered/specific -> use 'the'
                        if re.search(r"\d|[A-Za-z]\d|\d[A-Za-z]", txt):
                            return f"the {item_text.strip()}"
                        # otherwise generic -> use 'a'
                        return f"a {item_text.strip()}"

                    # 🧩 Merge "key and badge" smoothly if both exist
                    unique_items = sorted(set(item_list), key=item_list.index)
                    if "key" in unique_items and "badge" in unique_items:
                        item_list_str = "a key and a badge"
                    elif any(i.endswith("s") for i in unique_items):
                        item_list_str = " and ".join(unique_items)
                    else:
                        item_list_str = " and ".join(
                            smart_item_phrase(i.strip()) for i in unique_items if i.strip()
                        ) or "a key"

                    # Recipient / authorized / location detection
                    recipient_match = re.search(r"\bfor\s+([A-Za-z\s\-\(\)]+)", action, re.IGNORECASE)
                    authorized_match = re.search(r"\(authorized by\s*([A-Za-z\s]+)\)", action, re.IGNORECASE)
                    location_match = re.search(r"\(([^)]+)\)$", action)

                    recipient = recipient_match.group(1).strip() if recipient_match else ""
                    authorized = authorized_match.group(1).strip() if authorized_match else ""
                    location = location_match.group(1).strip() if location_match else ""

                    # 🧾 Build final polished sentence
                    text_parts = [
                        f"conducted key service and provided {item_list_str}",
                    ]
                    if recipient:
                        text_parts.append(f"to {recipient}")
                    if authorized:
                        text_parts.append(f"(authorized by {authorized})")
                    if location:
                        text_parts.append(f"at {location}")

                    text_parts.append("ensuring controlled access.")
                    buffer["action"] = " ".join(text_parts).strip()

                elif re.search(r"\b(returned|collected|retrieved|received back)\b", action, re.IGNORECASE):
                    # Extract possible key/badge identifiers
                    item_list = re.findall(r"\b(key\d*|badge\d*|key|badge|keys|badges)\b", action, re.IGNORECASE)
                    item_list = [i.lower() for i in item_list]

                    # 🧠 Smart article + plural logic (reuse same function)
                    def smart_item_phrase(item_text: str) -> str:
                        if not item_text:
                            return ""
                        txt = item_text.strip().lower()
                        if txt.endswith("s") and not re.search(r"\d", txt):
                            return item_text.strip()
                        if re.search(r"\d|[A-Za-z]\d|\d[A-Za-z]", txt):
                            return f"the {item_text.strip()}"
                        return f"a {item_text.strip()}"

                    # 🧩 Merge "key and badge" smoothly if both exist
                    unique_items = sorted(set(item_list), key=item_list.index)
                    if "key" in unique_items and "badge" in unique_items:
                        item_list_str = "a key and a badge"
                    elif any(i.endswith("s") for i in unique_items):
                        item_list_str = " and ".join(unique_items)
                    else:
                        item_list_str = " and ".join(
                            smart_item_phrase(i.strip()) for i in unique_items if i.strip()
                        ) or "a key"

                    # Recipient / authorized / location detection
                    recipient_match = re.search(r"\bfrom\s+([A-Za-z\s\-\(\)]+)", action, re.IGNORECASE)
                    authorized_match = re.search(r"\(authorized by\s*([A-Za-z\s]+)\)", action, re.IGNORECASE)
                    location_match = re.search(r"\(([^)]+)\)$", action)

                    recipient = recipient_match.group(1).strip() if recipient_match else ""
                    authorized = authorized_match.group(1).strip() if authorized_match else ""
                    location = location_match.group(1).strip() if location_match else ""

                    # 🧾 Build final polished sentence
                    text_parts = [
                        f"conducted key service and processed the return of {item_list_str}",
                    ]
                    if recipient:
                        text_parts.append(f"from {recipient}")
                    if authorized:
                        text_parts.append(f"(authorized by {authorized})")
                    if location:
                        text_parts.append(f"at {location}")

                    text_parts.append("confirming full accountability and reinventory.")
                    buffer["action"] = " ".join(text_parts).strip()

                # --- Default fallback (keep and polish original action narrative) ---
                else:
                    cleaned_action = action.strip().rstrip(".")
                    # Clean awkward trailing words
                    cleaned_action = re.sub(r"\bthe\s*$", "", cleaned_action, flags=re.IGNORECASE)

                    # Make the first letter lowercase if it starts mid-sentence (e.g., "Granted" → "granted")
                    cleaned_action = re.sub(r"^([A-Z])", lambda m: m.group(1).lower(), cleaned_action)

                    if re.search(r"\b(grant|escort|coordinate|assist|verify|monitor|support|supervise|respond)\b", cleaned_action, re.IGNORECASE):
                        buffer["action"] = (
                            f"conducted key service and {to_past_tense(cleaned_action)}. "
                            f"{'Ensured proper coor+dination and authorized access' if not re.search(r'ensure|authorized|access', cleaned_action, re.IGNORECASE) else ''}"
                        ).strip()
                    elif re.search(r"\b(access|entry|visit|contractor|vendor|staff)\b", cleaned_action, re.IGNORECASE):
                        buffer["action"] = (
                            f"conducted key service and {to_past_tense(cleaned_action)}. "
                            f"{'Verified authorization and maintained secure access control' if not re.search(r'verify|authorization|security', cleaned_action, re.IGNORECASE) else ''}"
                        ).strip()
                    else:
                        buffer["action"] = (
                            f"conducted key service and {to_past_tense(cleaned_action)}. "
                            f"Ensured safety and authorized access during operation."
                        ).strip()

                # 🧹 Final grammar cleanup: fix duplicate 'the the' or 'doors doors'
                buffer["action"] = re.sub(r"\bthe\s+the\b", "the", buffer["action"], flags=re.IGNORECASE)
                buffer["action"] = re.sub(r"\bdoors\s+doors\b", "doors", buffer["action"], flags=re.IGNORECASE)

                # --- Auto-format location name globally ---
                if buffer.get("location"):
                    # Clean and normalize location (e.g., "rooftop" → "Rooftop")
                    formatted_location = format_location_name(buffer["location"])
                    
                    # Replace buffer["location"] with HTML-tagged version
                    buffer["location"] = f"Location: <b>{formatted_location}</b>"


                # --- Build event line ---
                evt = build_event_line(buffer)

                # 🧹 Post-clean in case noise reappears from merged buffer fields
                if evt:
                    evt = clean_shift_noise(evt)
                    parsed[sec].append(evt)

        # LOADING DOCK: similar polish to Key Service
        elif sec == "Loading Dock Access (Lock & Unlock)":
            
            action = (buffer.get("action") or "").strip()
            company = (buffer.get("company") or "").strip()
    
            if action:
                action = to_past_tense(action)

                # Unlock scenarios
                if "unlock" in action.lower() or "open" in action.lower():
                    buffer["action"] = (
                        f"was dispatched to the loading dock in response to delivery needs and unlocked the gate for {company}, "
                        f"granting authorized access and facilitating scheduled operations."
                    )

                # Lock / Secure scenarios
                elif "lock" in action.lower() or "secure" in action.lower() or "close" in action.lower():
                    buffer["action"] = (
                        f"was dispatched to the loading dock and secured the gate after {company}'s delivery, "
                        f"maintaining site safety and compliance."
                    )

                # Default fallback
                else:
                    buffer["action"] = (
                        f"was dispatched to the loading dock and {action} for {company}"
                    )

            # Build the event line
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)

        # FIRE PANEL: compliance-oriented handling
        elif sec == "Fire Panel Bypass/Online":
            action = (buffer.get("action") or "").strip()
            company = (buffer.get("company") or "").strip()
            lower = action.lower()

            # --- Detect hold types ---
            hold_types = []
            if "full" in lower:
                hold_types.append("full hold")
            if "supervisory" in lower:
                hold_types.append("supervisory hold")
            if "trouble" in lower:
                hold_types.append("trouble hold")

            # Default: if nothing specific, assume supervisory
            if not hold_types and ("hold" in lower or "extend" in lower or "bypass" in lower or "put" in lower):
                hold_types = ["supervisory hold"]

            # Merge supervisory + trouble into a combined string
            if "supervisory" in lower and "trouble" in lower:
                hold_type_str = "supervisory and trouble hold"
            else:
                hold_type_str = " and ".join(hold_types) if hold_types else "system hold"

            # --- Extract explicit time (normalize formats) ---
            time_match = re.search(
                r'(\d{1,2}:\d{2}\s*[AP]M|\d{3,4}\s*[AP]M|\d{1,2}\s*[AP]M)',
                action,
                re.IGNORECASE
            )
            hold_until = None
            if time_match:
                raw_time = time_match.group(1).upper().replace(" ", "")
                if re.match(r'^\d{3,4}[AP]M$', raw_time):  # e.g., 0200PM
                    digits = re.sub(r'[AP]M', '', raw_time)
                    ampm = "AM" if "A" in raw_time else "PM"
                    if len(digits) == 3:  # e.g., 200PM
                        digits = "0" + digits
                    hh, mm = digits[:2], digits[2:]
                    hold_until = f"{hh}:{mm} {ampm}"
                elif re.match(r'^\d{1,2}[AP]M$', raw_time):  # e.g., 2PM
                    hh = raw_time[:-2].zfill(2)
                    hold_until = f"{hh}:00 {raw_time[-2:]}"
                else:
                    hold_until = raw_time

            # --- Build polished action text ---
            if "extend" in lower or "extended" in lower:
                buffer["action"] = (
                    f"conducted fire panel operations and extended the {hold_type_str}"
                    + (f" until {hold_until}" if hold_until else " until the scheduled time")
                    + f" in coordination with {company or 'the vendor'}"
                )

            elif "hold" in lower or "bypass" in lower or "put" in lower or "place" in lower:
                buffer["action"] = (
                    f"conducted fire panel operations and put the system on {hold_type_str}"
                    + (f" until {hold_until}" if hold_until else " until the scheduled time")
                    + f" in coordination with {company or 'the vendor'}"
                )

            elif any(x in lower for x in ["restore", "restored", "back online", "bring online", "brought online", "online", "remove", "removed"]):
                if "full" in lower:
                    buffer["action"] = (
                        f"conducted fire panel operations and restored the system from full hold "
                        f"in coordination with {company or 'the vendor'}"
                    )
                elif "supervisory" in lower and "trouble" in lower:
                    buffer["action"] = (
                        f"conducted fire panel operations and restored the system from supervisory and trouble hold "
                        f"in coordination with {company or 'the vendor'}"
                    )
                elif "supervisory" in lower:
                    buffer["action"] = (
                        f"conducted fire panel operations and restored the system from supervisory hold "
                        f"in coordination with {company or 'the vendor'}"
                    )
                elif "trouble" in lower:
                    buffer["action"] = (
                        f"conducted fire panel operations and restored the system from trouble hold "
                        f"in coordination with {company or 'the vendor'}"
                    )
                else:
                    buffer["action"] = (
                        f"conducted fire panel operations and restored the system online "
                        f"in coordination with {company or 'the vendor'}"
                    )

            else:
                buffer["action"] = (
                    f"conducted fire panel operations and extended the {hold_type_str}"
                    + (f" until {hold_until}" if hold_until else " until the scheduled time")
                    + f" in coordination with {company or 'the vendor'}"
                )

            # --- Append final event line ---
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)

        # AES PHONE CALLS: professional, compliance-oriented handling
        elif sec == "AES Phone Calls":
            action = (buffer.get("action") or "").strip()
            company = (buffer.get("company") or "").strip()
            operator_name = (buffer.get("operator_name") or "N/A").strip()
            operator_number = (buffer.get("operator_number") or "N/A").strip()

            lower = action.lower()

            # --- Detect hold types ---
            has_supervisory = "supervisory" in lower
            has_trouble = "trouble" in lower
            has_full = "full" in lower

            if has_full:
                hold_type_str = "full hold"
            elif has_supervisory and has_trouble:
                hold_type_str = "supervisory and trouble hold"
            elif has_supervisory:
                hold_type_str = "supervisory hold"
            elif has_trouble:
                hold_type_str = "trouble hold"
            elif any(x in lower for x in ["hold", "extend", "test"]):
                hold_type_str = "supervisory hold"
            else:
                hold_type_str = "system hold"

            # --- Extract explicit time (normalize formats) ---
            time_match = re.search(
                r'(\d{1,2}:\d{2}\s*[AP]M|\d{3,4}\s*[AP]M|\d{1,2}\s*[AP]M)',
                action,
                re.IGNORECASE
            )
            hold_until = None
            if time_match:
                raw_time = time_match.group(1).upper().replace(" ", "")
                if re.match(r'^\d{3,4}[AP]M$', raw_time):  # e.g., 0200PM
                    digits = re.sub(r'[AP]M', '', raw_time)
                    ampm = "AM" if "A" in raw_time else "PM"
                    if len(digits) == 3:
                        digits = "0" + digits
                    hh, mm = digits[:2], digits[2:]
                    hold_until = f"{hh}:{mm} {ampm}"
                elif re.match(r'^\d{1,2}[AP]M$', raw_time):  # e.g., 2PM
                    hh = raw_time[:-2].zfill(2)
                    hold_until = f"{hh}:00 {raw_time[-2:]}"
                else:
                    hold_until = raw_time

            # --- Build polished AES call text ---
            if "extend" in lower:
                buffer["action"] = (
                    f"called the AES Alarm Monitoring and extended the {hold_type_str}"
                    + (f" until {hold_until}" if hold_until else "")
                    + f" in coordination with {company or 'the vendor'} "
                    + f"<font color='green'>(Operator Name: <b>{operator_name}</b>, Operator Number: <b>{operator_number}</b>)</font>"
                )

            elif any(x in lower for x in ["hold", "bypass", "test"]):
                buffer["action"] = (
                    f"called the AES Alarm Monitoring and placed the system on {hold_type_str}"
                    + (f" until {hold_until}" if hold_until else "")
                    + f" in coordination with {company or 'the vendor'} "
                    + f"<font color='green'>(Operator Name: <b>{operator_name}</b>, Operator Number: <b>{operator_number}</b>)</font>"
                )

            else:
                # Default catch-all
                buffer["action"] = (
                    f"called the AES Alarm Monitoring and placed the system on {hold_type_str}"
                    + (f" until {hold_until}" if hold_until else "")
                    + f" in coordination with {company or 'the vendor'} "
                    + f"<font color='green'>(Operator Name: {operator_name}, Operator Number: {operator_number})</font>"
                )

            # --- Append final event line ---
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)

        # JANITORIAL: professional & dynamic handling
        elif sec == "Janitorial":
            action = (buffer.get("action") or "").strip()
            company = (buffer.get("company") or "ABM Janitorial").strip()  # Default to ABM
            lower = action.lower()

            # --- NEW: Seattle Ambassadors dispatch handling ---
            if any(k in lower for k in ["seattle ambassadors", "ambassador", "mid call", "mid dispatch"]):
                buffer["action"] = (
                    "placed a phone call to MID to dispatch the Seattle Ambassadors on site "
                    "to clean human waste, bodily fluids, and messy trash on the exterior."
                )

            # --- Keyword-based categorization ---
            if any(k in lower for k in ["spill", "liquid", "water leak", "slip", "hazard"]):
                buffer["action"] = (
                    f"coordinated janitorial response and notified {company} to clean a reported spill/hazard "
                    f"to ensure safety and prevent accidents"
                )

            elif any(k in lower for k in ["trash", "garbage", "overflow", "waste", "dumpster"]):
                buffer["action"] = (
                    f"reported janitorial concern of trash overflow and dispatched {company} "
                    f"to clear the waste and maintain cleanliness"
                )

            elif any(k in lower for k in ["restroom", "toilet", "bathroom", "urinal", "supply", "paper towel", "soap"]):
                buffer["action"] = (
                    f"notified {company} regarding restroom cleaning and supply replenishment "
                    f"to maintain sanitary conditions"
                )

            elif any(k in lower for k in ["vacuum", "sweep", "mop", "sanitize", "disinfect"]):
                buffer["action"] = (
                    f"assigned {company} to perform floor care and sanitization tasks, "
                    f"including vacuuming, mopping, or sweeping as required"
                )

            elif any(k in lower for k in ["odor", "smell", "stain", "debris", "dirty", "cleaning required"]):
                buffer["action"] = (
                    f"requested {company} to address reported odor, stains, or debris "
                    f"to restore a clean and professional environment"
                )

            else:
                # Fallback when no keyword detected
                buffer["action"] = (
                    f"coordinated janitorial services through {company} "
                    f"to address reported cleaning needs on site"
                )

            # --- Auto-format location name globally ---
            if buffer.get("location"):
                buffer["location"] = format_location_name(buffer["location"])

            # --- Append final event line ---
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)

        # --- Other/Miscellaneous → Additional Information ---
        # --- SPECIAL CASE: Other/Miscellaneous → Additional Information ---
        elif sec == "Additional Information":
            n = len(lines)
            i = 0
            while i < n:
                ln = lines[i]
                # print("DEBUG: searching all lines for 4:17 AM...")
                # for idx, l in enumerate(lines):
                #     if "4:17" in l or "04:17" in l:
                #         print(idx, "|", repr(l))
                # 🧩 Detect "Other / Miscellaneous" header OR timestamps in same section
                if (
                    re.search(r"\bOther\s*/?\s*Miscellaneous\b", ln, re.I)
                    or (
                        parsed.get("Additional Information")
                        and not re.search(r"END\s*OF\s*REPORT|DAILY\s*ACTIVITY|^Page\s+\d+", ln, re.I)
                        and (
                            re.match(r"^\d{1,2}/\d{1,2}/\d{2,4}\s+\d{1,2}:\d{2}", ln)
                            or re.match(r"^Start\s*(?:Date|Time)\s*:", ln, re.I)
                        )
                    )
                ):



                    # 🧹 Skip fake hits (like footer fragments)
                    if re.search(r"REPORT\s*-\s*LOGBOOK|Generated\s+on", ln, re.I):
                        i += 1
                        continue
                    misc_buffer = {"category": "Additional Information"}

                    # 🧩 1️⃣ Officer name (above or nearby)
                    officer_name = ""

                    # 🔼 Look upward for officer or supervisor lines
                    for j in range(i - 1, max(0, i - 8), -1):
                        t = lines[j].strip()

                        if not t or re.search(r"(geolocation|comment|start\s*date|end\s*date|report|multi-line|tour|actually|needed)", t, re.I):
                            continue

                        # ✅ Match any pattern like:
                        #  - GETACHEW Tegegne
                        #  - GETACHEW Tegegne (Officers)
                        #  - GETACHEW TEGEGNE (Site Supervisors)
                        m_name = re.match(
                            r"^([A-Z][A-Za-z]+)\s+([A-Z][A-Za-z]+)(?:\s*\((?:Officers?|Site\s*Supervisors?)\))?$",
                            t,
                            re.I,
                        )
                        if m_name:
                            first_part, second_part = m_name.groups()

                            # Detect and fix uppercase "LAST FIRST" → "First Last"
                            if first_part.isupper() and (second_part[0].isupper() and second_part[1:].islower() or second_part.isupper()):
                                # Reverse order: LAST FIRST → First Last
                                officer_name = f"{second_part.capitalize()} {first_part.capitalize()}"
                            else:
                                # Keep order: First Last
                                officer_name = f"{first_part.capitalize()} {second_part.capitalize()}"
                            break


                    # 🔽 Fallback: look downward if not found (for cases like 'PAYMAN Ramazan' after NEW ACTIVITY)
                    if not officer_name:
                        for j in range(i + 1, min(n, i + 10)):
                            t = lines[j].strip()
                            if re.match(r"^\s*$", t):
                                continue
                            if re.match(r"^\s*(?:NEW\s+ACTIVITY|COMMENTS?)", t, re.I):
                                continue
                            m_down = re.match(r"^([A-Z][A-Za-z]+)\s+([A-Z][A-Za-z]+)(?:\s*\(Officers?\))?$", t)
                            if m_down:
                                first, last = m_down.groups()
                                officer_name = f"{first.capitalize()} {last.capitalize()}"
                                break

                    if officer_name:
                        misc_buffer["officer"] = officer_name

                    # 🧩 2️⃣ Gather this block (improved to handle all entries)
                    block_lines = []
                    next_idx = n
                    for k in range(i, n):
                        line_k = lines[k].strip()

                        # 🧱 Stop when we clearly reach the end of the Additional Information section
                        if re.search(r"^(End\s*of\s*Report|Daily\s*Activity|Summary\s*of|Work\s*Orders|Patrol\s*Check|Log\s*Summary)", line_k, re.I):
                            next_idx = k
                            break

                        # 🧹 skip only real footer lines (do NOT skip timestamps)
                        if re.search(r"(REPORT\s*-\s*LOGBOOK|Generated\s+on)", line_k, re.I):
                            continue

                        # 🧭 Detect the end of this block
                        if k > i:
                            # ✅ break on a new "Start Date" (beginning of a new entry)
                            if re.match(r"^Start\s*(?:Date|Time)\s*:\s*\d{1,2}/\d{1,2}/\d{2,4}\s+\d{1,2}:\d{2}", line_k, re.I):
                                next_idx = k
                                break
                            # ✅ break on a plain timestamp line
                            if re.match(r"^\d{1,2}/\d{1,2}/\d{2,4}\s+\d{1,2}:\d{2}", line_k):
                                next_idx = k
                                break
                            # ✅ break on "NEW ACTIVITY" just in case
                            if re.match(r"^\s*NEW\s+ACTIVITY\b", line_k, re.I):
                                next_idx = k
                                break

                            # Or if a new "Other/Miscellaneous" appears far enough apart (not same page)
                            if (
                                re.match(r"^\s*Other\s*/?\s*Miscellaneous\b", line_k, re.I)
                                and (k - i) > 3  # ensure it's not the same paragraph
                                and not re.search(r"REPORT\s*-\s*LOGBOOK|Generated\s+on", line_k, re.I)
                            ):
                                next_idx = k
                                break

                            # Otherwise keep collecting lines normally
                        block_lines.append(line_k)

                    # 🧩 3️⃣ Extract Start Date (priority: line before “Comments”)
                    start_date = ""
                    for idx, line in enumerate(block_lines):
                        if re.search(r"comments?", line, re.I) and idx > 0:
                            prev_line = block_lines[idx - 1]
                            m_prev = re.search(
                                r"Start\s*(?:Date|Time)\s*:\s*([0-9/]+\s+\d{1,2}:\d{2}\s*(?:AM|PM)?)",
                                prev_line,
                                re.I,
                            )
                            if m_prev:
                                start_date = m_prev.group(1).strip()
                                break

                    # fallback — look inside block if not found
                    if not start_date:
                        m_any = re.search(
                            r"Start\s*(?:Date|Time)\s*:\s*([0-9/]+\s+\d{1,2}:\d{2}\s*(?:AM|PM)?)",
                            "\n".join(block_lines),
                            re.I,
                        )
                        if m_any:
                            start_date = m_any.group(1).strip()

                    # fallback — look upward for timestamp
                    if not start_date:
                        for j in range(i - 1, max(0, i - 6), -1):
                            m_time = re.search(r"([0-9/]+\s+\d{1,2}:\d{2}\s*(?:AM|PM)?)", lines[j])
                            if m_time:
                                start_date = m_time.group(1).strip()
                                break
                    misc_buffer["start_date"] = start_date

                    # 🧩 4️⃣ Extract full multi-line comment safely
                    block_text = "\n".join(block_lines)

                    # Find where the "Multi-line text field :" starts
                    m_comment_start = None
                    for idx, line in enumerate(block_lines):
                        if re.search(r"Multi-line\s*text\s*field\s*:", line, re.I):
                            m_comment_start = idx
                            break

                    comment_lines = []
                    if m_comment_start is not None:
                        for k in range(m_comment_start, len(block_lines)):
                            line_k = block_lines[k].strip()
                            # 🛑 Stop at next NEW ACTIVITY or TOUR or another Misc header
                            if re.match(r"^\s*(NEW\s+ACTIVITY|TOUR\s|Other\s*/?\s*Miscellaneous)\b", line_k, re.I):
                                break

                            # 🧹 Skip footer/page fragments
                            if re.search(r"(REPORT\s*-\s*LOGBOOK|Generated\s+on|^Page\s+\d+|^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2})", line_k, re.I):
                                continue

                            # Remove the "Multi-line text field :" label from first line
                            line_k = re.sub(r"^-?\s*Multi-line\s*text\s*field\s*:\s*", "", line_k, flags=re.I).strip()

                            if line_k:
                                comment_lines.append(line_k)

                    comment = " ".join(comment_lines).strip()

                    # 🧹 Clean up noise
                    comment = re.sub(r"\b[Cc]lose\b", "", comment).strip().rstrip(".")
                    comment = re.sub(r"\s*\(.*?\)", "", comment).strip()
                    comment = re.sub(r"\s+", " ", comment).strip()
                    if comment and not comment.endswith("."):
                        comment += "."

                    # If no comment found, skip this block cleanly
                    if not comment:
                        i = next_idx
                        continue

                    # 🏷️ 5️⃣ Smart location inference
                    location = ""
                    text = comment.lower()

                    non_location_entities = [
                        "kone", "davis", "fedex", "usps", "abm", "cedar", "brunson", "engineer", "technician"
                    ]

                    # 1️⃣ Escort / delivery pattern (e.g., "Escorted Kone to floor 9")
                    m_move = re.search(
                        r"\b(?:escorted|accompanied|guided|took|brought|delivered|walked|assisted)\s+\w+(?:\s+\w+)?\s+to\s+([A-Za-z0-9\s\-]+?)(?:[.,;]|$)",
                        text, re.I
                    )
                    if m_move:
                        candidate = m_move.group(1).strip(" .,-")
                        # 🚫 Skip single-letter, pronouns, verbs, or known names
                        if (
                            len(candidate) == 1
                            or re.match(r"^(i|me|my|we|they|he|she|you)$", candidate.lower())
                            or re.match(r"^(escort|escorte|assist|help|report|inform|notify|contact)$", candidate.lower())
                            or any(candidate.lower().startswith(x) for x in non_location_entities)
                        ):
                            candidate = ""
                        if candidate:
                            location = candidate

                    # 2️⃣ Fallback: general preposition pattern (at/in/on/etc.)
                    if not location:
                        m_loc1 = re.search(
                            r"\b(?:at|in|on|inside|near|around|to)\s+([A-Za-z0-9\-\s]+?)(?:[.,;]|$)", text, re.I
                        )
                        if m_loc1:
                            candidate = m_loc1.group(1).strip(" .,-")
                            # Skip junk like "I", "me", or verbs
                            if not re.match(r"^(i|me|my|we|they|he|she|you)$", candidate.lower()):
                                location = candidate

                    # 3️⃣ Fallback: if still nothing but "floor" present
                    if not location and "floor" in text:
                        m_floor = re.search(r"(?:on|to)\s+(?:the\s+)?([A-Za-z0-9\s]*floor\s*\d*)", text)
                        if m_floor:
                            location = m_floor.group(1).strip()

                    # 🧹 4️⃣ Final cleanup and validation
                    if location:
                        # skip single letters or pronouns again, just in case
                        if len(location) == 1 or location.lower() in ["i", "me", "my", "we", "they", "he", "she", "you"]:
                            location = ""
                        else:
                            location = re.sub(r"\s+", " ", location).strip()
                            # Capitalize properly
                            parts = location.split()
                            location = " ".join([p.capitalize() if not re.match(r"^[A-Z0-9]+$", p) else p for p in parts])

                    # 🩹 Fallback: if still empty, default to "N/A" (strict uppercase)
                    if not location:
                        location = "N/A"

                    # 🧩 Clean up & beautify
                    if location and location != "N/A":
                        location = re.sub(r"\s+", " ", location).strip()
                        # Fix small common abbreviations or names
                        replacements = {
                            "sb": "SB",
                            "nb": "NB",
                            "eb": "EB",
                            "wb": "WB",
                            "p1": "P1",
                            "l1": "L1",
                            "subbasement": "Subbasement",
                            "loading dock": "Loading Dock",
                            "rooftop": "Rooftop",
                        }
                        for k, v in replacements.items():
                            if location.lower() == k:
                                location = v
                                break

                        # Capitalize gracefully (multi-word case)
                        parts = location.split()
                        location = " ".join(
                            [p.capitalize() if not re.match(r"^[A-Z0-9]+$", p) else p for p in parts]
                        )

                    # 🕕 6️⃣ Format and build event
                    date_fmt = _fmt_date_for_line(start_date)
                    officer = misc_buffer.get("officer", "").strip()

                    # ✅ Skip invalid trailing lines (no date or no comment)
                    if not date_fmt or not comment:
                        i += 1
                        continue

                    evt = f"{date_fmt} – {bold_officer(officer)} has reported {comment}"

                    # ✅ Always show location — inferred or N/A
                    final_location = location if location else "N/A"
                    evt += f" (<font color='red'>Location: <b>{final_location}</b></font>)"

                    # 🧾 7️⃣ Append if unique
                    parsed.setdefault("Additional Information", [])
                    if evt not in parsed["Additional Information"]:
                        parsed["Additional Information"].append(evt)

                    # ✅ Always continue scanning from the next timestamp boundary
                    if next_idx > i:
                        i = next_idx           # jump directly to the start of the next entry
                    else:
                        i += 1                 # safety step
                    continue



                i += 1

        else:
            # Normal path for every other section
            evt = build_event_line(buffer)
            if evt:
                parsed[sec].append(evt)
                # Count transient either by classification OR by the tag flag (extra safety)
                if sec == "Transient Removal" or transient_tag_seen:
                    transient_count += 1

        # Reset buffers/flags for next event
        buffer = {}
        labels.clear()
        transient_tag_seen = False
        last_field = None
        
    
    # for ln in lines:
    for i, ln in enumerate(lines):
        if not ln:
            continue
        # --- Prevent multi-line 'All persons involved...' continuation from merging into previous field ---
        # Sometimes a wrapped line starts with "numbers) :", which belongs to the next field
        if re.match(r"^\s*numbers\)\s*:\s*", ln, re.IGNORECASE):
            last_field = None
            continue

        # --- Capture bare timestamp lines (EXCLUSIVE to Incident Report) ---
        if "incident report" in buffer.get("category", "").lower():
            # Match "Start Date : 9/30/2025 3:47 AM"
            m_start = re.search(r"start\s*date\s*:\s*([0-9/]+\s+\d{1,2}:\d{2}\s*[APap][Mm])", ln, re.IGNORECASE)
            if m_start:
                start_val = m_start.group(1).strip()
                if buffer.get("date"):
                    flush_event()
                    buffer.clear()
                buffer["start_date"] = start_val
                buffer["date"] = start_val  # use this as the event timestamp
                last_field = "date"
                continue
        
        # --- Capture bare timestamp lines (EXCLUSIVE to Elevator Entrapment Incident) ---
        if "elevator entrapment incident" in buffer.get("category", "").lower():
            # Match "Start Date : 9/30/2025 4:30 AM"
            m_start = re.search(
                r"start\s*date\s*:\s*([0-9/]+\s+\d{1,2}:\d{2}\s*[APap][Mm])",
                ln,
                re.IGNORECASE
            )
            if m_start:
                start_val = m_start.group(1).strip()
                if buffer.get("date"):
                    flush_event()
                    buffer.clear()
                buffer["start_date"] = start_val
                buffer["date"] = start_val  # use this as the event timestamp
                last_field = "date"
                continue

        
        # --- Capture "Who Called" in Escalation section ---
        if re.search(r"(if\s*so,?\s*who\s*called|who\s*called\s*them)", ln, re.IGNORECASE):
            val = ln.split(":", 1)[-1].strip()

            # ✅ Collect continuation lines until "Vehicle Information" or another field
            idx = lines.index(ln)
            for nxt in lines[idx + 1:]:
                s = nxt.strip()
                # Stop at next major marker or Vehicle Information
                if re.match(r"^(vehicle|synopsis|report|incident|description|time|location|escalation)\b", s, re.IGNORECASE):
                    break
                if s:
                    val += " " + s

            # ✅ Clean up common formatting & artifacts
            val = re.sub(r"^\W+", "", val).strip()
            val = re.sub(r"\s{2,}", " ", val)
            val = re.sub(r"([a-z])([A-Z])", r"\1 \2", val)  # fix mashed words like "AliAhmed"
            val = val.rstrip(")").strip()

            if val:
                buffer["who_called"] = val
            continue

        # --- Capture Parties Involved (multi-line, comma-separated, includes same-line value) ---
        if re.search(r"all\s*persons\s*involved", ln, re.IGNORECASE):
            val_lines = []

            # --- capture text after colon, even if wrapped to next line ---
            val = ln.split(":", 1)[-1].strip()
            idx = lines.index(ln)

            # If nothing after the colon, check the very next line
            if not val and idx + 1 < len(lines):
                nxt_line = lines[idx + 1].strip()
                # only treat it as continuation if not a new field header
                if not re.match(r"^(-\s*description|description|vehicle|report|escalation|incident|synopsis)\b", nxt_line, re.IGNORECASE):
                    val = nxt_line

            if val:
                val_lines.append(val)

            # --- Collect continuation lines until "- Description" or another header ---
            for nxt in lines[idx + 1:]:
                s = nxt.strip()
                if re.match(r"^(-\s*description|description|vehicle|report|escalation|incident|synopsis)\b", s, re.IGNORECASE):
                    break
                if not s:
                    continue
                if re.search(r"REPORT\s*-\s*LOGBOOK\s*PDF", s, re.IGNORECASE):
                    continue
                if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}", s):  # footer timestamps
                    continue
                # avoid double-adding if it was already captured as next-line of the colon
                if s != val:
                    val_lines.append(s)

            # ✅ Join all lines with commas
            combined = ", ".join(val_lines)

            # ✅ Clean label, fix spacing and mashups
            combined = re.sub(r"^-?\s*All\s*persons\s*involved.*?:", "", combined, flags=re.IGNORECASE)
            combined = re.sub(r"\s{2,}", " ", combined)
            combined = re.sub(r"([a-z])([A-Z])", r"\1 \2", combined)
            combined = combined.strip(" -,:").strip()

            # ✅ Add “and” before the last item — unless it already contains “and”
            if combined and "," in combined:
                parts = [p.strip() for p in combined.split(",") if p.strip()]
                if len(parts) > 1:
                    last = parts[-1]
                    # only add “and” if it’s not already in the final chunk
                    if not re.search(r"\band\b", last, re.IGNORECASE):
                        combined = ", ".join(parts[:-1]) + ", and " + last
                    else:
                        combined = ", ".join(parts)
                else:
                    combined = parts[0]

            if combined:
                buffer["parties_involved"] = combined
            continue

        # --- Capture bare timestamp lines (EXCLUSIVE to Key Service) ---
        if "key service" in buffer.get("category", "").lower():
            if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*[APap][Mm]$", ln):
                # Flush only when Key Service already has a timestamp (isolating multiple entries)
                if buffer.get("date"):
                    flush_event()
                    buffer.clear()
                buffer["date"] = ln.strip()
                last_field = "date"
                continue
        # --- Capture bare timestamp lines (EXCLUSIVE to Loading Dock) ---
        if "loading dock" in buffer.get("category", "").lower() or "dock gate" in buffer.get("category", "").lower():
            if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*[APap][Mm]$", ln):
                # if we already built a valid entry, flush it before starting the next
                if buffer.get("date") and (buffer.get("action") or buffer.get("company")):
                    flush_event()
                    buffer.clear()
                buffer["date"] = ln.strip()
                last_field = "date"
                continue
                
        # --- Capture bare timestamp lines (EXCLUSIVE to Fire Panel) ---
        if "fire panel" in buffer.get("category", "").lower():
            # Only match true timestamps, skip inline 'until/by/at' times
            if (
                re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*[APap][Mm]$", ln)
                and not re.search(r"\b(until|by|at)\b", ln, re.IGNORECASE)
            ):
                # If we already have a valid event, flush before starting a new one
                if buffer.get("date") and (buffer.get("action") or buffer.get("company")):
                    flush_event()
                    buffer.clear()

                buffer["date"] = ln.strip()
                last_field = "date"
                continue

        # 🧩 Detect officer name immediately above "Loading Dock Gate"
        # Supported formats:
        #   1️⃣ TEGEGNE Getachew
        #   2️⃣ TEGEGNE Getachew (Officers)
        #   3️⃣ TEGEGNE GETACHEW (Site Supervisors)
        if re.search(r"\bLoading\s+Dock\s+Gate\b", ln, re.I):
            # Look a few lines above for the officer name
            for back in range(i - 1, max(0, i - 6), -1):
                prev_line = lines[back].strip()

                # Match full-name patterns and optional role suffix
                # e.g., "TEGEGNE Getachew", "TEGEGNE Getachew (Officers)", "TEGEGNE GETACHEW (Site Supervisors)"
                if re.match(r"^[A-Z][A-Za-z]+\s+[A-Z][A-Za-z]+(?:\s*\((?:Officers?|Site\s+Supervisors?)\))?$", prev_line, re.I):
                    # Clean up any parentheses text like "(Officers)" or "(Site Supervisors)"
                    name = re.sub(r"\(.*?\)", "", prev_line).strip()
                    parts = name.split()

                    # 🧠 Detect LAST FIRST format (first word uppercase, second capitalized)
                    # Swap to First Last → "Getachew Tegegne"
                    if len(parts) == 2:
                        # If both are uppercase or mixed like TEGEGNE GETACHEW or TEGEGNE Getachew
                        if parts[0].isupper() or parts[0][0].isupper():
                            buffer["officer"] = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                        else:
                            # Otherwise, keep as-is and normalize capitalization
                            buffer["officer"] = " ".join(p.capitalize() for p in parts)
                    else:
                        # Handle multi-part names (e.g., MOHMAND Faiz Mohammad)
                        buffer["officer"] = " ".join(p.capitalize() for p in parts)

                    break
        
        # 🧩 Detect officer name immediately above "Fire Panel Bypass/Online"
        # Handles:
        #   1️⃣ TEGEGNE Getachew
        #   2️⃣ TEGEGNE Getachew (Officers)
        #   3️⃣ TEGEGNE GETACHEW (Site Supervisors)
        #   4️⃣ MOHMAND Faiz Mohammad → Faiz Mohammad ✅
        if re.search(r"\bFire\s+Panel\s+Bypass/Online\b", ln, re.I):
            for back in range(i - 1, max(0, i - 6), -1):
                prev_line = lines[back].strip()

                # 🚫 Skip irrelevant or noisy lines
                if re.search(r"^(new activity|300 pine|close|start date|geolocation|page|report)", prev_line, re.I):
                    continue

                # ✅ Match officer names (with or without role suffix)
                if re.match(r"^[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,2}(?:\s*\((?:Officers?|Site\s+Supervisors?)\))?$", prev_line, re.I):
                    # Clean "(Officers)" or "(Site Supervisors)"
                    name = re.sub(r"\(.*?\)", "", prev_line).strip()
                    parts = name.split()

                    # 🔄 Normalize LAST FIRST → First Last
                    if len(parts) == 2 and parts[0].isupper():
                        # Example: TEGEGNE Getachew → Getachew Tegegne
                        buffer["officer"] = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                    elif len(parts) == 3 and parts[0].isupper():
                        # Example: MOHMAND Faiz Mohammad → Faiz Mohammad
                        buffer["officer"] = f"{parts[1].capitalize()} {parts[2].capitalize()}"
                    else:
                        # Example: Already normal (Faiz Mohammad or Rafiullah Zakaria)
                        buffer["officer"] = " ".join(p.capitalize() for p in parts)
                    break
                
        # 🧩 Detect officer name immediately above "AES Phone Call"
        # Handles:
        #   1️⃣ TEGEGNE Getachew
        #   2️⃣ TEGEGNE Getachew (Officers)
        #   3️⃣ TEGEGNE GETACHEW (Site Supervisors)
        #   4️⃣ MOHMAND Faiz Mohammad → Faiz Mohammad ✅
        if re.search(r"\bAES\s+Phone\s+Call\b", ln, re.I):
            for back in range(i - 1, max(0, i - 6), -1):
                prev_line = lines[back].strip()

                # 🚫 Skip irrelevant or noisy lines
                if re.search(r"^(new activity|300 pine|close|start date|geolocation|page|report)", prev_line, re.I):
                    continue

                # ✅ Match officer names (with or without role suffix)
                if re.match(r"^[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,2}(?:\s*\((?:Officers?|Site\s+Supervisors?)\))?$", prev_line, re.I):
                    # Clean "(Officers)" or "(Site Supervisors)"
                    name = re.sub(r"\(.*?\)", "", prev_line).strip()
                    parts = name.split()

                    # 🔄 Normalize LAST FIRST → First Last
                    if len(parts) == 2 and parts[0].isupper():
                        # Example: TEGEGNE Getachew → Getachew Tegegne
                        buffer["officer"] = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                    elif len(parts) == 3 and parts[0].isupper():
                        # Example: MOHMAND Faiz Mohammad → Faiz Mohammad
                        buffer["officer"] = f"{parts[1].capitalize()} {parts[2].capitalize()}"
                    else:
                        # Example: Already normal (Faiz Mohammad or Rafiullah Zakaria)
                        buffer["officer"] = " ".join(p.capitalize() for p in parts)
                    break

        # --- SPD New Activity detection (split line case) ---
        # if ln.strip().lower().startswith("new activity"):
        #     next_idx = lines.index(ln) + 1
        #     for j in range(next_idx, min(next_idx + 6, len(lines))):
        #         maybe_officer = lines[j].strip()
        #         # Detect pattern like "ZAKARIA Rafiullah (Officers)"
        #         if re.match(r"^[A-Z][A-Za-z]+\s+[A-Z][A-Za-z]+", maybe_officer) and "(officer" in maybe_officer.lower():
        #             # normalize "ZAKARIA Rafiullah (Officers)" → "Rafiullah Zakaria"
        #             name = re.sub(r"\(.*?\)", "", maybe_officer).strip()
        #             parts = name.split()
        #             if len(parts) == 2 and parts[0].isupper():
        #                 name = f"{parts[1].capitalize()} {parts[0].capitalize()}"
        #             buffer["officer"] = name
        #             break

        # --- SPD New Activity detection (split line case) ---
        if ln.strip().lower().startswith("new activity"):
            next_idx = lines.index(ln) + 1
            for j in range(next_idx, min(next_idx + 6, len(lines))):
                maybe_officer = lines[j].strip()
                if re.match(r"^[A-Z][A-Za-z]+\s+[A-Z][A-Za-z]+", maybe_officer):
                    name = re.sub(r"\(.*?\)", "", maybe_officer).strip()

                    # 🧩 Normalize officer name (supports Officers / Site Supervisors / plain)
                    name = re.sub(r"\s*\((?:Officers?|Site\s+Supervisors?)\)\s*", "", name, flags=re.I).strip()

                    m_last_first = re.match(r"^([A-Z][A-Za-z]+)\s+([A-Z][A-Za-z]+)$", name)
                    if m_last_first:
                        last, first = m_last_first.groups()
                        if first.islower() or first.istitle():
                            name = f"{first.capitalize()} {last.capitalize()}"
                        else:
                            name = f"{first.capitalize()} {last.capitalize()}"
                    elif re.match(r"^[A-Z]{2,}\s+[A-Z]{2,}$", name):
                        parts = name.split()
                        if len(parts) == 2:
                            name = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                    elif re.match(r"^[A-Z][a-z]+\s+[A-Z][a-z]+$", name):
                        first, last = name.split()[:2]
                        name = f"{first.capitalize()} {last.capitalize()}"

                    buffer["officer"] = re.sub(r"\s{2,}", " ", name).strip()
                    break

        
        # --- SPD Presence / Emergency Response on Site ---
        if "spd presence/emergency response on" in ln.lower():
            in_spd = True
            spd_buffer = {"category": "SPD Presence/Emergency Response on Site"}

            # 🚫 Clear any old carried start_date
            buffer.pop("start_date", None)

            # Carry officer name from above
            if buffer.get("officer"):
                spd_buffer["officer"] = buffer["officer"].strip()

            # ✅ Look ahead strictly within next 3 lines for "Start Date"
            cur_idx = lines.index(ln)
            for nxt_ln in lines[cur_idx : cur_idx + 4]:
                m_start = re.search(r"start\s*date\s*:\s*([\d/]+\s+\d{1,2}:\d{2}\s*(?:[APap][Mm])?)", nxt_ln, re.I)
                if m_start:
                    start_val = m_start.group(1).strip()
                    spd_buffer["start_date"] = start_val
                    buffer["start_date"] = start_val
                    break

            continue  # go to SPD block handling next iteration

        if in_spd:
            # --- Start Date (always capture fresh within this SPD block) ---
            if re.search(r"start\s*date\s*:", ln, re.IGNORECASE):
                start_val = ln.split(":", 1)[-1].strip()
                spd_buffer["start_date"] = start_val
                # also store in buffer to reuse in same block scope
                buffer["start_date"] = start_val


            # --- Officer (rare inside block, but keep if present) ---
            elif re.search(r"-\s*officer\s*:", ln, re.IGNORECASE):
                spd_buffer["officer"] = ln.split(":", 1)[-1].strip()

            # --- Incident Date / Time ---
            elif re.search(r"date\s*of\s*incident\s*:", ln, re.IGNORECASE):
                spd_buffer["incident_date"] = ln.split(":", 1)[-1].strip()

            elif re.search(r"time\s*of\s*incident\s*:", ln, re.IGNORECASE):
                raw_t = ln.split(":", 1)[-1].strip()
                m = re.match(r"(\d{1,2}):(\d{2})", raw_t)
                if m:
                    hh, mm = map(int, m.groups())
                    ampm = "AM"
                    if hh >= 12:
                        ampm = "PM"
                        if hh > 12:
                            hh -= 12
                    elif hh == 0:
                        hh = 12
                    spd_buffer["incident_time"] = f"{hh}:{mm:02d} {ampm}"
                else:
                    spd_buffer["incident_time"] = raw_t.upper().replace("HRS", "").strip()

            # --- Location (ignore geolocation coordinates) ---
            elif ln.lower().startswith("- location"):
                loc = ln.split(":", 1)[-1].strip()
                if "°" not in loc:
                    spd_buffer["location"] = format_location_name(loc)

            # --- Who called SPD? ---
            elif re.search(r"who\s+called\s+spd", ln, re.IGNORECASE):
                caller_parts = []

                # Get everything after the colon on the same line
                val = ln.split(":", 1)[-1].strip()
                if val:
                    caller_parts.append(val)

                idx = lines.index(ln)

                # Collect continuation lines (usually next 1–2 lines)
                for nxt in lines[idx + 1:]:
                    s = nxt.strip()
                    # stop before next field
                    if re.match(r"^-*\s*(parties|images|upload picture|date|time|location)\b", s, re.IGNORECASE):
                        break
                    if not s or s.lower().startswith("parties involved"):
                        break
                    if re.search(r"REPORT\s*-\s*LOGBOOK\s*PDF", s, re.IGNORECASE):
                        continue
                    if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}", s):
                        continue
                    caller_parts.append(s)

                # ✅ Join without commas (since it’s usually a person’s name)
                caller = " ".join(caller_parts)
                caller = re.sub(r"(?i)\bsecurity\s+officer\b", "", caller).strip()
                caller = re.sub(r"\s+", " ", caller).strip()
                caller = caller.title()  # Normalize capitalization like “Mohamed Mohamed”

                spd_buffer["caller"] = caller


            # --- Parties Involved (multi-line, comma-separated + 'and') ---
            elif re.search(r"parties\s*involved", ln, re.IGNORECASE):
                val_lines = []

                # capture after colon (same line)
                val = ln.split(":", 1)[-1].strip()
                if val:
                    val_lines.append(val)

                idx = lines.index(ln)
                for nxt in lines[idx + 1:]:
                    s = nxt.strip()
                    # ✅ stop *before* "Images" or "Upload picture" lines, with or without dashes
                    if re.match(r"^-*\s*(images|upload picture|date|time|location|who called)\b", s, re.IGNORECASE):
                        break
                    if not s:
                        continue
                    if re.search(r"REPORT\s*-\s*LOGBOOK\s*PDF", s, re.IGNORECASE):
                        continue
                    if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}", s):
                        continue
                    val_lines.append(s)

                # ✅ Join with commas
                combined = ", ".join(v.strip() for v in val_lines if v.strip())

                # ✅ Clean and normalize
                combined = re.sub(r"^-?\s*parties\s*involved.*?:", "", combined, flags=re.IGNORECASE)
                combined = re.sub(r"\s{2,}", " ", combined)
                combined = combined.strip(" -,:").strip()

                # ✅ Add “and” before the last entry if multiple
                if combined and "," in combined:
                    parts = [p.strip() for p in combined.split(",") if p.strip()]
                    if len(parts) > 1:
                        last = parts[-1]
                        if not re.search(r"\band\b", last, re.IGNORECASE):
                            combined = ", ".join(parts[:-1]) + ", and " + last
                        else:
                            combined = ", ".join(parts)
                    else:
                        combined = parts[0]

                spd_buffer["parties"] = combined

            # --- Long Description (multi-line) ---
            elif "long description of incident" in ln.lower():
                # collect narrative until next field header or end of block
                idx = lines.index(ln)
                desc = ln.split(":", 1)[-1].strip()
                for nxt in lines[idx + 1:]:
                    s = nxt.strip()
                    if re.match(r"^-+\s*(who called|parties|images|upload picture|date|time|location)\b", s, re.IGNORECASE):
                        break
                    if s:
                        desc += " " + s

                # tidy common OCR typos
                # ✅ Clean common OCR typos and grammatical noise
                desc = re.sub(r"\b5he\b", "the", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\baed\b", "A", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\bw\s*he\b", "when he", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\bn on\b", " on", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\s+", " ", desc).strip()
                # ensure first letter capitalized
                desc = desc[0].upper() + desc[1:] if desc else desc

                # ensure it explicitly includes transfer-to-hospital sentence once
                if not re.search(r"\btransferred to the hospital\b", desc, re.IGNORECASE):
                    if not desc.endswith("."):
                        desc += "."
                    desc += " The individual was transferred to the hospital."
                else:
                    if not desc.endswith("."):
                        desc += "."

                # ✅ Officer handled by build_event_line() — no "Officer … reported that" prefix here
                officer = (spd_buffer.get("officer") or "").strip()
                if officer:
                    officer_norm = re.sub(r"\s+", " ", officer).strip()
                    parts = officer_norm.split()
                    if len(parts) == 2 and parts[0].isupper():
                        officer_norm = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                    spd_buffer["officer"] = officer_norm

                # ✅ Ensure description starts cleanly and capitalized
                # ✅ Smart narrative for SPD report
                desc = desc.strip()

                # Tidy OCR typos again (safety)
                desc = re.sub(r"\baed\b", "A", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\bw\s*he\b", "when he", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\b5he\b", "the", desc, flags=re.IGNORECASE)
                desc = re.sub(r"\s+", " ", desc).strip()

                # Officer is already printed by build_event_line() – do NOT add it here again
                officer = (spd_buffer.get("officer") or "").strip()
                if officer:
                    officer_norm = re.sub(r"\s+", " ", officer).strip()
                    parts = officer_norm.split()
                    if len(parts) == 2 and parts[0].isupper():
                        officer_norm = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                    spd_buffer["officer"] = officer_norm

                # Sentence starts naturally, prefixed only with “reported that”
                # 🧠 Smart prefix for SPD narrative: lowercase the first natural word
                desc = re.sub(r"\s+", " ", desc.strip())

                if desc:
                    # Lowercase only if first token looks like a normal word (not acronym or number)
                    first_token = desc.split(" ", 1)[0]
                    if re.match(r"^[A-Z][a-z]?$", first_token):  # matches "A", "At", "Around"
                        desc = first_token.lower() + desc[len(first_token):]
                    # Always prefix with "reported that"
                    desc = f"reported that {desc}"
                else:
                    desc = "reported that an incident occurred on site."

                spd_buffer["action"] = desc

                # 🔹 Prevent double 'Officer Name' duplication in final output and
                # ✅ Prevent redundant "reported that" or repeated officer mentions
                if officer:
                    # remove any "officer <name>" phrases already inside description
                    desc = re.sub(rf"(?i)\bofficer\s+{re.escape(officer)}\b", "", desc).strip()
                    # remove any extra "reported that" duplication
                    desc = re.sub(r"(?i)\breported that\s+reported that\b", "reported that", desc)
                    # ensure it only starts with one clean "reported that"
                    if not desc.lower().startswith("reported that"):
                        desc = f"reported that {desc.strip()}"

                spd_buffer["action"] = desc

            # --- End of SPD block ---
            elif ln.strip().lower().startswith("new activity") or ln.strip().lower() == "close":
                idate = (spd_buffer.get("incident_date") or "").strip()
                itime = (spd_buffer.get("incident_time") or "").strip()
                
                if idate:
                    try:
                        mm, dd, yyyy = [int(x) for x in idate.split("/")]
                        spd_buffer["date"] = f"{mm:02d}/{dd:02d}/{str(yyyy)[-2:]} {itime}"
                    except:
                        spd_buffer["date"] = spd_buffer.get("start_date", "")
                else:
                    spd_buffer["date"] = spd_buffer.get("start_date", "")

                # start_val = (spd_buffer.get("start_date") or "").strip()
                # idate = (spd_buffer.get("incident_date") or "").strip()
                # itime = (spd_buffer.get("incident_time") or "").strip()

                # if start_val:
                #     # Already normalized above
                #     spd_buffer["date"] = start_val
                # elif idate:
                #     try:
                #         mm, dd, yyyy = [int(x) for x in idate.split("/")]
                #         spd_buffer["date"] = f"{mm:02d}/{dd:02d}/{str(yyyy)[-2:]} {itime}"
                #     except:
                #         spd_buffer["date"] = idate
                # else:
                #     spd_buffer["date"] = ""

                evt = build_event_line(spd_buffer)
                if evt:
                    evt = re.sub(r"–\s*\([^)]+\)\s*$", "", evt).strip()

                    info = []
                    if idate:
                        red = f"<font color='red'>Incident Date: <b>{idate}</b>"
                        if itime:
                            red += f" at <b>{itime}</b>"
                        red += "</font>"
                        info.append(red)
                    if spd_buffer.get("location"):
                        info.append(f"<font color='red'>Location: <b>{spd_buffer['location']}</b></font>")
                    if spd_buffer.get("caller"):
                        info.append(f"Who Called: <b>{spd_buffer['caller']}</b>")
                    if spd_buffer.get("parties"):
                        info.append(f"Parties Involved: <b>{spd_buffer['parties']}</b>")

                    if info:
                        evt += " (" + ", ".join(info) + ")"

                    parsed["SPD Presence/Emergency Response on Site"].append(evt)

                in_spd = False
                spd_buffer = {}
            continue
            

        # --- Detect start of a Seattle Ambassadors activity ---
        if ln.strip().lower() == "seattle ambassadors" or "seattle ambassadors start date" in ln.lower():
            # 🚫 Clear any carried start_date from previous sections
            buffer.pop("start_date", None)
            in_ambassador = True
            ambassador_buffer = {"category": "Seattle Ambassadors"}

            # ✅ Inline start date pattern (handles "Seattle Ambassadors Start Date : 9/29/2025 4:25 PM")
            m_inline = re.search(r"start\s*date\s*:\s*([\d/]+\s+\d{1,2}:\d{2}\s*(?:[APap][Mm])?)", ln, re.I)
            if m_inline:
                start_val = m_inline.group(1).strip()
                ambassador_buffer["start_date"] = start_val
                ambassador_buffer["date"] = start_val
                buffer["start_date"] = start_val
            continue

        # --- Capture data while inside Seattle Ambassadors block ---
        if in_ambassador:
            # ✅ Catch case where "Close" appears right before Seattle Ambassadors
            if ln.strip().lower() == "close":
                # Peek ahead — if next line starts with Seattle Ambassadors, skip this Close
                cur_idx = lines.index(ln)
                if cur_idx + 1 < len(lines) and "seattle ambassadors" in lines[cur_idx + 1].lower():
                    continue  # ignore this Close since new section starts next
                else:
                    # Normal Close behavior if not followed by Seattle Ambassadors
                    in_ambassador = False
                    ambassador_buffer = {}
                    continue

            # ✅ Start Date line on next line
            if re.search(r"start\s*date\s*:", ln, re.IGNORECASE):
                start_val = ln.split(":", 1)[-1].strip()
                ambassador_buffer["start_date"] = start_val
                buffer["start_date"] = start_val

                # Normalize 24hr or 12hr time to consistent format
                m = re.match(r"(\d{1,2})/(\d{1,2})/(\d{4})\s+(\d{1,2}):(\d{2})(?:\s*([APap][Mm]))?", start_val)
                if m:
                    mm, dd, yyyy, hh, mins, ampm = m.groups()
                    hh = int(hh)
                    if not ampm:  # 24-hour fix
                        ampm = "AM"
                        if hh >= 12:
                            ampm = "PM"
                            if hh > 12:
                                hh -= 12
                        elif hh == 0:
                            hh = 12
                    formatted = f"{int(mm):02d}/{int(dd):02d}/{yyyy[-2:]} {hh}:{mins} {ampm.upper()}"
                    ambassador_buffer["date"] = formatted
                else:
                    ambassador_buffer["date"] = start_val
                continue

            # ✅ Capture Time and store separately
            if re.search(r"-\s*time\s*:", ln, re.IGNORECASE):
                raw_time = ln.split(":", 1)[-1].strip()
                m = re.match(r"(\d{1,2}):(\d{2})", raw_time)
                if m:
                    hh, mm = map(int, m.groups())
                    ampm = "AM"
                    if hh >= 12:
                        ampm = "PM"
                        if hh > 12:
                            hh -= 12
                    elif hh == 0:
                        hh = 12
                    formatted_time = f"{hh}:{mm:02d} {ampm}"
                    ambassador_buffer["incident_time"] = formatted_time  # 👈 store here
                continue

            # Capture Incident Date (from "- Date :")
            elif re.search(r"-\s*date\s*:", ln, re.IGNORECASE):
                ambassador_buffer["incident_date"] = ln.split(":", 1)[-1].strip()
                continue

            # Officer
            if ln.lower().startswith("- officer") or "officer :" in ln.lower():
                ambassador_buffer["officer"] = ln.split(":", 1)[-1].strip()

            # Location
            elif ln.lower().startswith("- location"):
                ambassador_buffer["location"] = ln.split(":", 1)[-1].strip()

            # --- End of this activity block ---
            elif ln.strip().lower().startswith("new activity") or ln.strip().lower() == "close":
                # ✅ Use - Time : based time if available
                if ambassador_buffer.get("incident_date") and ambassador_buffer.get("incident_time"):
                    ambassador_buffer["date"] = f"{ambassador_buffer['incident_date']} {ambassador_buffer['incident_time']}"
                elif not ambassador_buffer.get("date") and ambassador_buffer.get("start_date"):
                    ambassador_buffer["date"] = ambassador_buffer["start_date"]

                # Clean and format location just once
                if ambassador_buffer.get("location"):
                    ambassador_buffer["location"] = format_location_name(
                        ambassador_buffer["location"].strip()
                    )

                # 🧠 Smart, clean action text (matches IR-style tone)
                ambassador_buffer["action"] = (
                    "placed a phone call to MID to dispatch the Seattle Ambassadors on site "
                    "to clean human waste, bodily fluids, and messy trash on the exterior."
                )

                # Build the event line – build_event_line() already appends (Location)
                evt = build_event_line(ambassador_buffer)
                if evt:
                    parsed["Janitorial"].append(evt)

                # Reset state
                in_ambassador = False
                ambassador_buffer = {}
            continue

        # --- BRAND NEW: Unsecure Door Parsing (Retail Issues) ---
        if "unsecure door" in ln.lower():
            # ✅ Start fresh every time
            in_unsecure = True
            unsecure_buffer = {
                "category": "Retail Issues",
                "start_date": "",
                "incident_date": "",
                "incident_time": "",
                "officer": "",
                "location": "",
                "action": "",
            }

            # --- Look ahead 12 lines max for Start Date + Details block ---
            cur_idx = lines.index(ln)
            block = lines[cur_idx : cur_idx + 15]  # capture near block

            for sub in block:
                # Capture Start Date (for fallback only)
                if re.search(r"start\s*date\s*:", sub, re.IGNORECASE):
                    m = re.search(r"([\d/]+\s+\d{1,2}:\d{2}\s*(?:[APap][Mm])?)", sub)
                    if m:
                        unsecure_buffer["start_date"] = m.group(1).strip()

                # Capture Details Date
                if re.search(r"-\s*date\s*:", sub, re.IGNORECASE):
                    m = re.search(r"([\d/]+)", sub)
                    if m:
                        unsecure_buffer["incident_date"] = m.group(1).strip()

                # Capture Details Time
                if re.search(r"-\s*time\s*:", sub, re.IGNORECASE):
                    m = re.search(r"(\d{1,2}):(\d{2})", sub)
                    if m:
                        hh, mm = map(int, m.groups())
                        ampm = "AM"
                        if hh >= 12:
                            ampm = "PM"
                            if hh > 12:
                                hh -= 12
                        elif hh == 0:
                            hh = 12
                        unsecure_buffer["incident_time"] = f"{hh}:{mm:02d} {ampm}"

                # Officer
                if "pre-defined list" in sub.lower() or re.search(r"-\s*officer\s*:", sub, re.IGNORECASE):
                    unsecure_buffer["officer"] = sub.split(":", 1)[-1].strip()

                # Location
                if re.search(r"-\s*location\s*:", sub, re.IGNORECASE):
                    loc = sub.split(":", 1)[-1].strip()
                    if "°" not in loc:
                        unsecure_buffer["location"] = format_location_name(loc)

            # ✅ Compute the final date using Details (if available)
            idate = unsecure_buffer.get("incident_date", "")
            itime = unsecure_buffer.get("incident_time", "")
            sdate = unsecure_buffer.get("start_date", "")

            if idate and itime:
                try:
                    mm, dd, yyyy = [int(x) for x in idate.split("/")]
                    unsecure_buffer["date"] = f"{mm:02d}/{dd:02d}/{str(yyyy)[-2:]} {itime}"
                except:
                    unsecure_buffer["date"] = f"{idate} {itime}"
            elif itime:
                unsecure_buffer["date"] = itime
            else:
                unsecure_buffer["date"] = sdate

            # ✅ Compose clean narrative
            unsecure_buffer["action"] = (
                f"reported that a door at {unsecure_buffer.get('location', 'the site')} "
                "was found unsecured and was properly secured upon discovery."
            )

            # ✅ Build and append
            evt = build_event_line(unsecure_buffer)
            if evt:
                parsed["Retail Issues"].append(evt)

            in_unsecure = False
            unsecure_buffer = {}
            continue

   
        # Ignore TOUR blocks entirely (for events)
        if ln.upper().startswith("TOUR"):
            in_tour = True
            continue
        if in_tour and (ln.startswith("NEW ACTIVITY") or ln.startswith("Start Date")):
            in_tour = False
        if in_tour:
            continue

        # Capture category/labels for classification
        if any(k in ln for k in [
            "AES Phone Call", "Loading Dock Gate", "Key Service", "Work Order", "Janitorial",
            "Incident Report", "Alarm", "Fire Panel Bypass/Online",
            "Transient Removal", "Retail", "Tenant", "Other/Miscellaneous",
            "Elevator Entrapment", "Entrapment Incident", "Stuck in Elevator"
        ]):
            
            # Special handling: Incident Report should flush as its own entry
            if "incident report" in ln.lower():
                carry = {k: buffer[k] for k in ("officer", "date") if k in buffer}
                flush_event()
                buffer.update(carry)

            # 👇 NEW: if we get another "Transient Removal" while one is already open, close the previous one
            if "transient removal" in ln.lower() and buffer.get("category", "").lower() == "transient removal":
                flush_event()

            labels.add(ln)
            buffer["category"] = ln.strip()
            if "Transient Removal" in ln:
                transient_tag_seen = True

        # Officer block (standard nested form)
        if ln == "Officer":
            waiting_officer_nested = True
            last_field = None
            continue
        if waiting_officer_nested and ln.startswith("- Officer :"):
            buffer["officer"] = ln.split(":", 1)[-1].strip()
            waiting_officer_nested = False
            last_field = "officer"
            continue

        if ln.startswith("- Officer :") or ln.startswith("Officer :"):
            buffer["officer"] = ln.split(":", 1)[-1].strip()
            last_field = "officer"
            continue

        # NEW: generic "(Officers)" line (works for Additional Info and anywhere else)
        m_off = OFFICER_LINE_RX.match(ln)
        if m_off:
            buffer["officer"] = m_off.group(1).strip(" -:")
            last_field = "officer"
            continue
        else:
            m_off_alt = OFFICER_LINE_RX_ALT.match(ln)
            if m_off_alt:
                buffer["officer"] = m_off_alt.group(1).strip(" -:")
                last_field = "officer"
                continue
        
        # 🧠 Normalize officer name: handle formatting and ordering for all officer entries
        # ---------------------------------------------------------------
        # This ensures officer names always appear in a consistent and human-friendly format.
        # Example conversions:
        #   "TEGEGNE Getachew"            → "Getachew Tegegne"
        #   "TEGEGNE GETACHEW (Site Supervisors)" → "Getachew Tegegne"
        #   "Zerihun Negussie"            → "Zerihun Negussie"
        #   "Loading Dock Gate"           → (ignored — not a person)
        # ---------------------------------------------------------------
        if buffer.get("officer"):
            name = buffer["officer"].strip()
            # Skip if looks like non-person label
            if not re.match(r"^(Activities|Loading|Key|Door|Gate|Victrola|Uniqlo|Report|Duration|Object)\b", name, re.I):
                parts = name.split()
                if len(parts) == 2 and parts[0].isupper() and parts[1][0].isupper():
                    # Swap order if looks like LAST FIRST
                    buffer["officer"] = f"{parts[1].capitalize()} {parts[0].capitalize()}"
                else:
                    # Otherwise just normalize capitalization
                    buffer["officer"] = " ".join(p.capitalize() for p in parts)

        # Flush NEW ACTIVITY only for Incident Reports
        if ln.startswith("NEW ACTIVITY") and buffer.get("category", "").lower() == "incident report":
            flush_event()
            continue
        
        # >>> ADD START: capture bare timestamp lines inside Incident Reports
        # if buffer.get("category", "").lower() == "incident report" and not buffer.get("date"):
        #     if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*[AP]M$", ln, re.IGNORECASE):
        #         buffer["date"] = ln.strip()
        #         last_field = "date"
        #         continue

        # New event starts
        if ln.startswith("Start Date"):
            if buffer.get("category", "").lower() == "incident report":
                # Inside an IR: this Start Date belongs to the current IR → do NOT flush here
                buffer["date"] = ln.split(":", 1)[-1].strip()
                last_field = "date"
                continue
            else:
                # Non-IR sections: Start Date begins a new event → flush previous
                flush_event()
                buffer["date"] = ln.split(":", 1)[-1].strip()
                last_field = "date"
                continue


        # Skip breaks flat-out
        if "Minute Break" in ln or "Lunch Break" in ln or "Break Details" in ln:
            last_field = None
            continue
        # Skip report-summary UI headers so they don't concatenate
        if any(x in ln for x in ("Totals Activities", "Total Activities", "Object Duration", "Activity Duration")):
            last_field = None
            continue

        # Skip stray "Close" markers (form UI artifacts)
        if ln.strip().lower() == "close":
            last_field = None
            continue
        # Skip Escalation / Evidence / callback headers so they don't pollute narratives
        if ln.lower().startswith("escalation?") or \
           ln.lower().startswith("- was the police") or \
           ln.lower().startswith("- (if so") or \
           ln.lower().startswith("- upload picture") or \
           ln.lower().startswith("- call back number") or \
           ln.lower().startswith("- all persons involved"):
            last_field = None
            continue

        # Content fields
        if "Security action" in ln or "What are you doing" in ln or "What did you do" in ln:
            buffer["action"] = ln.split(":", 1)[-1].strip()
            last_field = "action"
            continue

        # Capture incident description
        if ln.lower().startswith("- description of what happened"):
            buffer["incident_description"] = ln.split(":", 1)[-1].strip()
            last_field = "incident_description"
            continue

        if ln.lower().startswith("- operator name"):
            buffer["operator_name"] = ln.split(":", 1)[-1].strip()
            last_field = "operator_name"
            continue

        if ln.lower().startswith("- operator #") or ln.lower().startswith("- operator number"):
            buffer["operator_number"] = ln.split(":", 1)[-1].strip()
            last_field = "operator_number"
            continue


        # Comments belong only to the current incident
        if ln.lower().startswith("- comments"):
            buffer["incident_comments"] = ln.split(":", 1)[-1].strip()
            last_field = "incident_comments"
            continue

        # Case 1: Explicit "Start Date :" line
        # if ln.startswith("Start Date"):
        #     # Always use Start Date as the primary timestamp for officer line
        #     buffer["start_date"] = ln.split(":", 1)[-1].strip()
        #     buffer["date"] = buffer["start_date"]  # force overwrite
        #     last_field = "date"
        #     continue

        # Case 2: Bare timestamp line like "9/25/2025 1:03 PM"
        if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*[AP]M$", ln, re.IGNORECASE):
            if buffer.get("category", "").lower() == "incident report":
                if "start_date" not in buffer:
                    buffer["start_date"] = ln.strip()
                    buffer["date"] = buffer["start_date"]
                continue

        # 🕒 Extract Start Date (works for all categories: Incident, Elevator, etc.)
        if re.search(r"Start\s*Date\s*:", ln, re.IGNORECASE):
            m_start = re.search(r"Start\s*Date\s*:\s*([0-9/:\sAPMapm]+)", ln, re.IGNORECASE)
            if m_start:
                start_date_val = m_start.group(1).strip()
                buffer["start_date"] = start_date_val
                buffer["date"] = start_date_val   # always used for officer timestamp
                last_field = "date"
                continue
        # --- Capture Start Date (used for Elevator Entrapment and others) ---
        if re.search(r"\bStart\s*Date\s*:", ln, re.IGNORECASE):
            m_start = re.search(
                r"\bStart\s*Date\s*:\s*([0-9/]+\s+\d{1,2}:\d{2}\s*[APap][Mm])",
                ln,
                re.IGNORECASE
            )
            if m_start:
                start_val = m_start.group(1).strip()
                buffer["start_date"] = start_val
                buffer["date"] = start_val      # keeps timestamp consistent with other sections
                last_field = "start_date"
                continue

        # 📅 Extract Date of Incident or Incident Date (robust against variants)
        if re.search(r"(Date\s*of\s*(the\s*)?Incident|Incident\s*Date)\s*:", ln, re.IGNORECASE):
            m_incident_date = re.search(
                r"(Date\s*of\s*(the\s*)?Incident|Incident\s*Date)\s*:\s*([0-9/]+)",
                ln,
                re.IGNORECASE
            )
            if m_incident_date:
                buffer["incident_date"] = m_incident_date.group(3).strip()
                last_field = "incident_date"
                continue

        # ⏰ Extract Time of Incident or Incident Time (handles AM/PM)
        if re.search(r"(Time\s*of\s*(the\s*)?Incident|Incident\s*Time)\s*:", ln, re.IGNORECASE):
            m_incident_time = re.search(
                r"(Time\s*of\s*(the\s*)?Incident|Incident\s*Time)\s*:\s*([0-9:]+\s*(AM|PM|am|pm)?)",
                ln,
                re.IGNORECASE
            )
            if m_incident_time:
                buffer["incident_time"] = m_incident_time.group(3).strip()
                last_field = "incident_time"
                continue

        if ln.lower().startswith("- incident location"):
            buffer["location"] = ln.split(":", 1)[-1].strip()
            last_field = "location"
            continue
        # --- Capture Long Description of Incident ---
        if ln.lower().startswith("- long description of incident"):
            buffer["incident_description"] = ln.split(":", 1)[-1].strip()
            last_field = "incident_description"
            continue
        
        # ✅ Capture continuation lines for Long Description of Incident
        if last_field == "incident_description":
            stripped = ln.strip()

            # 1) stop if we hit “Parties Involved” (with or without dash)
            if re.match(r"(?i)^(-\s*)?parties\s+involved\b", stripped):
                # start capturing parties separately
                buffer["parties_involved"] = ""
                last_field = "parties_involved"
                continue

            # 2) stop on other known headers
            if re.match(r"(?i)^(-\s*)?(additional comments|photos?|evidence|incident info|geolocation)\b", stripped):
                last_field = None
                continue

            # 3) skip page/footer noise
            if re.search(r"(report\s*-\s*logbook\s*pdf|\bpage\s*\d+/\d+)", stripped, re.IGNORECASE):
                continue

            # otherwise, keep appending
            prev = buffer.get("incident_description", "")
            buffer["incident_description"] = (prev + " " + stripped).strip()
            continue

        # ✅ Capture "Parties Involved" block (multi-line, cleaned with commas + 'and' for Elevator Entrapment Incidents)
        if ln.strip().lower().startswith("parties involved"):
            val_lines = []

            # Get everything after colon on same line
            val = ln.split(":", 1)[-1].strip()
            if val:
                val_lines.append(val)

            idx = lines.index(ln)

            # Collect continuation lines until Photos/Evidence/Additional Comments
            for nxt in lines[idx + 1:]:
                s = nxt.strip()
                if re.match(r"(?i)^(-\s*)?(photos?|evidence|additional comments)\b", s):
                    break
                if not s:
                    continue
                if re.search(r"REPORT\s*-\s*LOGBOOK\s*PDF", s, re.IGNORECASE):
                    continue
                if re.match(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}", s):
                    continue
                val_lines.append(re.sub(r"^-\s*", "", s))  # remove leading dash

            # ✅ Join with commas
            combined = ", ".join(v.strip() for v in val_lines if v.strip())

            # ✅ Clean up: remove redundant label and double spaces
            combined = re.sub(r"^-?\s*parties\s*involved.*?:", "", combined, flags=re.IGNORECASE)
            combined = re.sub(r"\s{2,}", " ", combined)
            combined = re.sub(r"([a-z])([A-Z])", r"\1 \2", combined)
            combined = combined.strip(" -,:").strip()

            # ✅ Add "and" before the last entry (only if not already present)
            if combined and "," in combined:
                parts = [p.strip() for p in combined.split(",") if p.strip()]
                if len(parts) > 1:
                    last = parts[-1]
                    if not re.search(r"\band\b", last, re.IGNORECASE):
                        combined = ", ".join(parts[:-1]) + ", and " + last
                    else:
                        combined = ", ".join(parts)
                else:
                    combined = parts[0]

            # ✅ Store in buffer
            if combined:
                buffer["parties_involved"] = combined

            last_field = None
            continue

        # Capture vehicle info
        if ln.lower().startswith("- year"):
            buffer["year"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- make"):
            buffer["make"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- model"):
            buffer["model"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- color"):
            buffer["color"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- description"):
            buffer["vehicle_description"] = ln.split(":", 1)[-1].strip()
            continue


        # Capture the Comments block used by Other/Miscellaneous
        if ln.lower().startswith("comments"):
            last_field = "comment"
            continue
        m_ml = MULTILINE_RX.match(ln)
        if m_ml:
            buffer["comment"] = m_ml.group(1).strip()
            last_field = "comment"
            continue

        if ln in ("Details", "Call Details", "Date & Time", "Date/Time"):
            last_field = None
            continue

        if "Company" in ln or "Vendor" in ln:
            buffer["company"] = ln.split(":", 1)[-1].strip()
            last_field = "company"
            continue

        if ln.startswith("Location") or ln.startswith("- Location :"):
            parts = ln.split(":", 1)
            if len(parts) > 1:
                buffer["location"] = parts[1].strip()
                last_field = "location"
                continue

        if ln.startswith("Geolocation") or "LOGBOOK PDF" in ln:
            last_field = None
            continue

        # Vehicle info capture
        if ln.lower().startswith("- year"):
            buffer["year"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- make"):
            buffer["make"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- model"):
            buffer["model"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- color"):
            buffer["color"] = ln.split(":", 1)[-1].strip()
            continue
        if ln.lower().startswith("- description"):
            buffer["vehicle_description"] = ln.split(":", 1)[-1].strip()
            continue

        # Continuation lines for action/company/location/comment (not starters)
        if last_field in ("action", "company", "location", "comment") and not is_new_block_line(ln):
            prev = buffer.get(last_field, "")
            if len(prev) < 200:  # keep reasonable; prevents runaway concatenation
                buffer[last_field] = (prev + " " + ln).strip()
            continue
        
        # Continuations for incident multi-line fields
        if last_field in ("incident_description", "incident_comments") and not is_new_block_line(ln):
            prev = buffer.get(last_field, "")
            if len(prev) < 800:
                buffer[last_field] = (prev + " " + ln).strip()
            continue

        # Other lines are ignored

    # --- Final safety flush (prevents first/last entry loss) ---
    if buffer.get("date"):
        flush_event()


    # Ensure every section has at least "None to Report"
    for s in SECTIONS:
        if not parsed[s]:
            parsed[s] = ["None to Report"]

    # Replace Transient Removal with summary count (final step)
    if transient_count > 0:
        parsed["Transient Removal"] = [
            f'<font color="red">Within the last 24 hours (<b>{transient_count:02d}</b>), transients were removed from the property.</font>'
        ]
    else:
        parsed["Transient Removal"] = ["None to Report"]

    # --- NEW: sort each section by datetime (oldest → newest) ---
    for s in SECTIONS:
        if parsed[s] and parsed[s][0] != "None to Report":
            parsed[s].sort(key=_extract_dt)

    return parsed

def smart_item_phrase(item_text: str) -> str:
    """
    Adds natural 'the', 'a', or plural handling for key/badge phrases.
    Example:
      key  → 'a key'
      keys → 'keys'
      key1 → 'the key1'
      badge A → 'the badge A'
      badge → 'a badge'
    """
    if not item_text:
        return ""

    txt = item_text.strip().lower()

    # If plural, leave as-is (no 'the' or 'a')
    if txt.endswith("s") and not re.search(r"\d", txt):
        return item_text.strip()

    # If specific item (contains a number or code like key1, badge A2)
    if re.search(r"\d|[A-Za-z]\d|\d[A-Za-z]", txt):
        return f"the {item_text.strip()}"

    # Otherwise, generic singular item
    return f"a {item_text.strip()}"

def format_location_name(loc: str) -> str:
    """
    Smart location capitalization:
    - Capitalizes every main word
    - Keeps 'to', 'from', 'at', 'and', 'or', 'of', 'the', 'in' lowercase (unless first)
    - Preserves all-caps words like FCC, UNIQLO
    """
    import re
    if not loc:
        return ""
    loc = loc.strip()
    words = re.split(r"(\s+)", loc)
    skip_words = {"and", "or", "of", "the", "to", "from", "at", "in"}
    new_words = []
    for i, word in enumerate(words):
        if not word.strip():
            new_words.append(word)
            continue
        if word.isupper() and len(word) > 1:
            new_words.append(word)
            continue
        w = word.lower()
        if i == 0 or w not in skip_words:
            new_words.append(w.capitalize())
        else:
            new_words.append(w)
    return "".join(new_words)


def clean_shift_noise(text: str) -> str:
    """
    Remove embedded shift handover noise like:
    'Yes - new emails received during shift communicated to the next'
    Also cleans stray 'Yes'/'No' or leftover 'the,' fragments in actions.
    """
    if not text:
        return text

    t = text.replace("—", "-").replace("–", "-")
    t = re.sub(r"\s+", " ", t).strip()

    # --- Known noise phrases ---
    noise_patterns = [
        r"\b(yes|no)\s*-\s*new\s*emails\s*received\s*during\s*shift\s*communicated\s*to\s*the\s*next\??\s*:?\s*(yes|no)?",
        r"\bnew\s*emails\s*received\s*during\s*shift\s*communicated\s*to\s*the\s*next\??\s*:?\s*(yes|no)?",
        r"\bnew\s*work\s*orders\s*communicated\s*to\s*the\s*next\s*shift\??\s*:?\s*(yes|no)?",
        r"\bimportant\s*info\s*passed\s*down\s*for\s*the\s*shift\s*:?\s*(yes|no)?",
    ]
    for pat in noise_patterns:
        t = re.sub(pat, "", t, flags=re.IGNORECASE)

    # --- Clean up random Yes/No words ---
    t = re.sub(r"\bthe\s+yes\b", "the", t, flags=re.IGNORECASE)
    t = re.sub(r"\bthe\s+no\b", "the", t, flags=re.IGNORECASE)
    t = re.sub(r"\b(yes|no)[,.\s]+\b", "", t, flags=re.IGNORECASE)

    # --- Remove stray 'the,' fragments ---
    t = re.sub(r"\bthe\s*,\s*", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\bthe\s*\.\s*", "", t, flags=re.IGNORECASE)

    # --- Final polish ---
    t = re.sub(r"\s{2,}", " ", t)
    t = re.sub(r"\s*[-,:;]\s*$", "", t)
    return t.strip()


def build_event_line(buffer):
    """
    Build:  'MM/DD/YY HH:MM AM – <b>Officer Name</b> action [for Company] (Location)'
    """
    date = (buffer.get("date") or "").strip()
    m = re.search(r"(\d{1,2})/(\d{1,2})/(\d{4})\s+(\d{1,2}:\d{2}\s*[AP]M)", date, re.IGNORECASE)
    if m:
        mm, dd, yyyy, tm = m.groups()
        date = f"{int(mm):02d}/{int(dd):02d}/{yyyy[-2:]} {tm.upper()}"

    officer = bold_officer(buffer.get("officer", ""))

    action = (buffer.get("action") or "").strip()
    if action:
        action = to_past_tense(action)
        # Remove category residue like standalone 'Close'
        action = re.sub(r"\bClose\b", "", action).strip()
    company = (buffer.get("company") or "").strip()
    location = (buffer.get("location") or "").strip()

    # If action already contains the company name, don't duplicate 'for Company'
    if company and action and company.lower() in action.lower():
        company = ""

    parts = []
    if date:
        parts.append(date)
    if officer:
        parts.append(officer)
    if action:
        parts.append(action + (f" for {company}" if company else ""))
    if location:
        parts.append(f"({location})")

    return " – ".join(p for p in parts if p).strip(" –")


def generate_pdf(parsed_data, date_range_header, out_path):
    # --- Footer function (runs on each page) ---
    def draw_footer(canvas, doc):
        canvas.saveState()

        generated_on = datetime.now().strftime("%B %d, %Y %I:%M %p")

        # --- Gray footer banner ---
        footer_text = (
            "========================================= "
            "Created by Mirwais Shamsi"
            " ========================================"
        )

        # --- Text pieces ---
        page_text = f"Page {doc.page}"
        generated_text = f" | Generated on {generated_on}"

        # Coordinates
        x_center = letter[0] / 2.0
        y_position = 0.5 * inch

        # --- Line 1: gray banner ---
        canvas.setFont("Helvetica", 8)
        canvas.setFillColorRGB(0.4, 0.4, 0.4)
        canvas.drawCentredString(x_center, y_position, footer_text)

        # --- Line 2: measure widths for proper alignment ---
        canvas.setFont("Helvetica-Bold", 8)
        page_width = canvas.stringWidth(page_text, "Helvetica-Bold", 8)

        canvas.setFont("Helvetica", 8)
        gen_width = canvas.stringWidth(generated_text, "Helvetica", 8)

        total_width = page_width + gen_width

        # Compute starting X so that the combined text is centered
        start_x = x_center - total_width / 2

        # --- Draw "Page X" (black, bold) ---
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColorRGB(0, 0, 0)
        canvas.drawString(start_x, y_position - 10, page_text)

        # --- Draw " | Generated on ..." (gray, normal) ---
        canvas.setFont("Helvetica", 8)
        canvas.setFillColorRGB(0.4, 0.4, 0.4)
        canvas.drawString(start_x + page_width, y_position - 10, generated_text)

        canvas.restoreState()

    # --- PDF Setup ---
    doc = SimpleDocTemplate(
        out_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()
    title_center  = ParagraphStyle("title_center",  parent=styles["Title"],  alignment=TA_CENTER)
    normal_center = ParagraphStyle("normal_center", parent=styles["Normal"], alignment=TA_CENTER)

    story = []

    # --- Logo ---
    if os.path.exists(LOGO_FILE):
        story.append(Image(LOGO_FILE, width=250, height=120))
        story.append(Spacer(1, 12))

    # --- Title + Date Range ---
    story.append(Paragraph("<u>300 Pine Daily Report</u>", title_center))
    story.append(Paragraph(f"Date Range: {date_range_header}", normal_center))
    story.append(Spacer(1, 18))

    # --- Sections ---
    for section in SECTIONS:
        story.append(Paragraph(f"<b>{section}</b>", styles["Heading3"]))
        entries = parsed_data.get(section, [])

        if entries:
            for i, line in enumerate(entries):
                story.append(Paragraph(f"- {line}", styles["Normal"]))

                # 🔹 Add blank line ONLY for specific sections — except after last entry
                if (
                    section in ["Incident Reports (IR) / Alarms", "Elevator Entrapment Incidents", "SPD Presence/Emergency Response on Site"]
                    and i < len(entries) - 1
                ):
                    story.append(Spacer(1, 8))
        else:
            story.append(Paragraph("None to Report", styles["Normal"]))

        story.append(Spacer(1, 12))  # normal section spacer

    # --- Build PDF ---
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)

    print(f"✅ PDF Summary generated: {out_path}")
    print(f"📅 Detected Date Range: {date_range_header}")

def add_page_number(run):
    """Insert a real dynamic PAGE field code into a docx run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_paragraph_with_html(doc, text):
    """Convert basic HTML (<b>, <font color>, etc.) to Word formatting."""
    # Remove leading dash or bullet-like prefixes before parsing
    text = re.sub(r"^\s*[-•]\s*", "", text)
    # Split text into tokens preserving tags
    parts = re.split(r"(<[^>]+>)", text)
    p = doc.add_paragraph(style="List Bullet")

    current_color = None
    bold_active = False

    for part in parts:
        if not part:
            continue

        # --- handle tags ---
        if part.lower().startswith("<b>"):
            bold_active = True
            continue
        elif part.lower().startswith("</b>"):
            bold_active = False
            continue
        elif part.lower().startswith("<font"):
            m = re.search(r"color=['\"]?(#[0-9A-Fa-f]{3,6}|red|blue|green|black)['\"]?", part)
            if m:
                color_val = m.group(1).lower()
                color_map = {
                    "red": RGBColor(200, 0, 0),
                    "blue": RGBColor(0, 0, 200),
                    "green": RGBColor(0, 150, 0),
                    "black": RGBColor(0, 0, 0),
                }
                current_color = color_map.get(color_val, RGBColor(200, 0, 0))
            continue
        elif part.lower().startswith("</font"):
            current_color = None
            continue

        # --- handle normal text ---
        run = p.add_run(part)
        if bold_active:
            run.bold = True
        if current_color:
            run.font.color.rgb = current_color
        run.font.size = Pt(11)

    p.paragraph_format.space_after = Pt(4)
    return p


def generate_docx(parsed_data, date_range_header, out_path):
    """Generate editable .docx version of the DAR summary with formatting and footer."""
    doc = Document()

    # --- Add Logo (centered, if available) ---
    if os.path.exists(LOGO_FILE):
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(LOGO_FILE, width=Inches(3.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")  # spacer

    # --- Title Section ---
    title = doc.add_heading("300 Pine Daily Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f"Date Range: {date_range_header}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")  # spacer

    # --- Each Section ---
    for section in SECTIONS:
        doc.add_heading(section, level=2)
        entries = parsed_data.get(section, [])

        if entries and entries[0] != "None to Report":
            for i, line in enumerate(entries):
                add_paragraph_with_html(doc, line)  # ✅ removed the "- " prefix here

                # 🔹 Add a blank line ONLY for IR, Elevator, and SPD sections — except after last entry
                if (
                    section in ["Incident Reports (IR) / Alarms", "Elevator Entrapment Incidents", "SPD Presence/Emergency Response on Site"]
                    and i < len(entries) - 1
                ):
                    doc.add_paragraph("")  # visual separation between entries
        else:
            # 🔹 Add bullet for "None to Report"
            doc.add_paragraph("None to Report", style="List Bullet")

        doc.add_paragraph("")  # normal section spacer

    # --- Footer (each page) ---
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_on = datetime.now().strftime("%B %d, %Y %I:%M %p")
    footer_text = (
        "================================= "
        "Created by Mirwais Shamsi"
        " ================================"
    )

    run1 = footer_para.add_run(f"{footer_text}\n")
    run1.font.size = Pt(9)
    run1.font.color.rgb = RGBColor(100, 100, 100)
    run1.italic = True

    run2 = footer_para.add_run("Page ")
    add_page_number(run2)
    run2 = footer_para.add_run(" | Generated on " + generated_on)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(120, 120, 120)
    run2.italic = True

    # --- Save ---
    doc.save(out_path)
    print(f"✅ DOCX Summary generated: {out_path}")
    print(f"📅 Detected Date Range: {date_range_header}")


if __name__ == "__main__":
    # 1️⃣ Read all text lines from report
    if not os.path.exists(INPUT_FILE):
        raise FileNotFoundError(f"Input file not found: {INPUT_FILE}")

    lines = extract_text_lines(INPUT_FILE)

    # 2️⃣ Date range (header + filename token)
    date_range_header, token = parse_date_range(lines, INPUT_FILE)

    # 3️⃣ Parse events into sections
    parsed = parse_events(lines)

    # 4️⃣ Choose export format
    EXPORT_MODE = "pdf"  # change to "pdf/docx" for PDF output

    # 5️⃣ Create output path
    if EXPORT_MODE == "pdf":
        out_name = f"DAR Report-Output_{token}.pdf" if token and token != "unknown" else "DAR Report-Output.pdf"
    else:
        out_name = f"DAR Report-Output_{token}.docx" if token and token != "unknown" else "DAR Report-Output.docx"

    OUT_FILE = os.path.join(BASE_DIR, out_name)

    # 6️⃣ Generate chosen report
    if EXPORT_MODE == "pdf":
        generate_pdf(parsed, date_range_header, OUT_FILE)
    else:
        generate_docx(parsed, date_range_header, OUT_FILE)



